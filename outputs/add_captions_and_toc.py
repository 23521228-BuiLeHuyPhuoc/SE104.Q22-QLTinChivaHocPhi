import os
import re
import zipfile
from copy import deepcopy

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


INPUT = "Nhom5QLTCVTHP.docx"
OUTPUT = os.path.join("outputs", "Nhom5QLTCVTHP_numbered.docx")


FIGURE_TITLES = [
    "Sơ đồ luồng dữ liệu lập danh sách năm học",
    "Sơ đồ luồng dữ liệu lập danh sách học kỳ",
    "Sơ đồ luồng dữ liệu lập danh sách tỉnh",
    "Sơ đồ luồng dữ liệu lập danh sách phường/xã",
    "Sơ đồ luồng dữ liệu lập danh sách khoa",
    "Sơ đồ luồng dữ liệu lập danh sách ngành học",
    "Sơ đồ luồng dữ liệu lập danh sách dân tộc",
    "Sơ đồ luồng dữ liệu lập danh sách đối tượng ưu tiên",
    "Sơ đồ luồng dữ liệu nhập thông tin hồ sơ sinh viên",
    "Sơ đồ luồng dữ liệu nhập danh sách môn học",
    "Sơ đồ luồng dữ liệu lập danh sách điều kiện môn học",
    "Sơ đồ luồng dữ liệu lập chương trình khung",
    "Sơ đồ luồng dữ liệu nhập chương trình học",
    "Sơ đồ luồng dữ liệu lập danh sách môn học mở",
    "Sơ đồ luồng dữ liệu lập danh sách điểm tổng kết học phần",
    "Sơ đồ luồng dữ liệu lập phiếu đăng ký học phần",
    "Sơ đồ luồng dữ liệu lập phiếu thu học phí",
    "Sơ đồ luồng dữ liệu lập danh sách sinh viên chưa đóng học phí",
    "Sơ đồ luồng dữ liệu tra cứu sinh viên",
    "Kiến trúc ba lớp của hệ thống",
    "Sơ đồ logic bảng NAMHOC",
    "Sơ đồ logic ràng buộc năm học và tham số",
    "Sơ đồ logic bảng TINH",
    "Sơ đồ logic ràng buộc tỉnh và tham số",
    "Sơ đồ logic bảng PHUONGXA và TINH",
    "Sơ đồ logic bảng KHOA",
    "Sơ đồ logic ràng buộc khoa và tham số",
    "Sơ đồ logic bảng NGANH",
    "Sơ đồ logic quan hệ KHOA, NGANH và tham số",
    "Sơ đồ logic bảng DANTOC",
    "Sơ đồ logic ràng buộc dân tộc và tham số",
    "Sơ đồ logic bảng DOITUONG",
    "Sơ đồ logic đối tượng ưu tiên và tham số",
    "Sơ đồ logic hồ sơ sinh viên",
    "Sơ đồ logic sinh viên và tham số hệ thống",
    "Sơ đồ logic môn học và tham số",
    "Sơ đồ logic điều kiện môn học",
    "Sơ đồ logic chương trình khung",
    "Sơ đồ logic chương trình học",
    "Sơ đồ logic môn học mở",
    "Sơ đồ logic đăng ký học phần và phiếu đăng ký",
    "Sơ đồ logic tham số học phí",
    "Sơ đồ logic đơn giá tín chỉ và học kỳ",
    "Sơ đồ logic báo cáo sinh viên nợ học phí",
    "Sơ đồ logic báo cáo doanh thu",
    "Sơ đồ dữ liệu phân quyền người dùng",
    "Sơ đồ logic dữ liệu hoàn chỉnh",
]


SPECIAL_TABLE_TITLES = {
    2: "Danh mục viết tắt",
    3: "Danh sách yêu cầu phần mềm",
    20: "Quy định học phí và đăng ký học phần",
    27: "Mô tả các thành phần trong kiến trúc hệ thống",
    28: "Danh sách bảng dữ liệu trong sơ đồ logic",
    65: "Danh sách các màn hình trong hệ thống",
    154: "Cài đặt và thử nghiệm các chức năng",
    155: "Phân công công việc nhóm",
}


def text_of_paragraph_element(el):
    return "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()


def normalize_spaces(text):
    return re.sub(r"\s+", " ", text or "").strip()


def set_run_font(run, size=14, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), "Times New Roman")


def ensure_caption_style(doc, style_name):
    styles = doc.styles
    if style_name in styles:
        style = styles[style_name]
    else:
        style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    style.font.bold = False
    style.font.italic = False
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style.paragraph_format.space_before = Pt(3)
    style.paragraph_format.space_after = Pt(6)
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), "Times New Roman")
    return style


def remove_paragraph(paragraph):
    el = paragraph._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def remove_generated_captions(doc):
    for paragraph in list(doc.paragraphs):
        style_name = paragraph.style.name if paragraph.style is not None else ""
        if style_name in {"CaptionHinh", "CaptionBang"}:
            remove_paragraph(paragraph)


def body_elements(doc):
    return list(doc.element.body)


def find_paragraph_element(doc, exact_text):
    for el in body_elements(doc):
        if el.tag == qn("w:p") and text_of_paragraph_element(el) == exact_text:
            return el
    raise ValueError(f"Không tìm thấy đoạn: {exact_text}")


def clear_between_headings(doc, start_text, end_text):
    elems = body_elements(doc)
    start = end = None
    for idx, el in enumerate(elems):
        if el.tag == qn("w:p"):
            txt = text_of_paragraph_element(el)
            if txt == start_text:
                start = idx
            elif start is not None and txt == end_text:
                end = idx
                break
    if start is None or end is None:
        raise ValueError(f"Không tìm thấy vùng giữa {start_text!r} và {end_text!r}")
    for el in elems[start + 1 : end]:
        el.getparent().remove(el)


def insert_paragraph_after(el, doc, text="", style=None, alignment=None, bold=False):
    new_p = OxmlElement("w:p")
    el.addnext(new_p)
    paragraph = Paragraph(new_p, doc._body)
    if style is not None:
        paragraph.style = style
    if alignment is not None:
        paragraph.alignment = alignment
    if text:
        run = paragraph.add_run(text)
        set_run_font(run, 14, bold=bold)
    return paragraph


def insert_paragraph_before(el, doc, text="", style=None, alignment=None, bold=False):
    new_p = OxmlElement("w:p")
    el.addprevious(new_p)
    paragraph = Paragraph(new_p, doc._body)
    if style is not None:
        paragraph.style = style
    if alignment is not None:
        paragraph.alignment = alignment
    if text:
        run = paragraph.add_run(text)
        set_run_font(run, 14, bold=bold)
    return paragraph


def reset_front_matter_lists(doc):
    clear_between_headings(doc, "MỤC LỤC", "DANH MỤC HÌNH ẢNH")
    clear_between_headings(doc, "DANH MỤC HÌNH ẢNH", "DANH MỤC BẢNG")
    clear_between_headings(doc, "DANH MỤC BẢNG", "DANH MỤC VIẾT TẮT")

    for heading_text, token in [
        ("MỤC LỤC", "__TOC_MAIN__"),
        ("DANH MỤC HÌNH ẢNH", "__TOC_FIGURES__"),
        ("DANH MỤC BẢNG", "__TOC_TABLES__"),
    ]:
        heading_el = find_paragraph_element(doc, heading_text)
        insert_paragraph_after(
            heading_el,
            doc,
            token,
            style=doc.styles["Normal"],
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
        )


def is_blank_table(table):
    return not normalize_spaces(" ".join(cell.text for row in table.rows for cell in row.cells))


def section_name_from_heading(text):
    text = normalize_spaces(text).rstrip(":")
    text = re.sub(r"^\d+(?:\.\d+)*\.?\s*", "", text)
    return text


def table_title(table_index, table, section, prev_text):
    if table_index in SPECIAL_TABLE_TITLES:
        return SPECIAL_TABLE_TITLES[table_index]

    section_clean = normalize_spaces(section).rstrip(":")
    prev_clean = normalize_spaces(prev_text).rstrip(":")
    header = " | ".join(cell.text.strip() for cell in table.rows[0].cells) if table.rows else ""

    match = re.match(r"^4\.4\.\d+\s+Bảng\s+(.+)", section_clean)
    if match:
        return f"Mô tả thuộc tính bảng {match.group(1).rstrip(':')}"

    screen_match = re.match(r"^5\.3\.\d+\.\s+(.+)", section_clean)
    if screen_match:
        screen = screen_match.group(1).rstrip(":")
        screen_lower = screen[:1].lower() + screen[1:]
        if "Biến cố" in header or "Xử lý" in header:
            return f"Danh sách biến cố xử lý trên {screen_lower}"
        return f"Mô tả đối tượng trên {screen_lower}"

    if table_index in range(4, 27):
        if prev_clean:
            return f"Biểu mẫu {prev_clean[:1].lower() + prev_clean[1:]}"
        return f"Biểu mẫu nghiệp vụ {table_index}"

    if prev_clean:
        prev_clean = re.sub(r"^[-•]\s*", "", prev_clean)
        return prev_clean

    if section_clean:
        return section_name_from_heading(section_clean)

    return f"Bảng dữ liệu {table_index}"


def add_table_captions(doc):
    table_by_element = {id(table._element): table for table in doc.tables}
    caption_no = 0
    original_table_idx = 0
    current_heading = ""
    previous_text = ""
    inserted = []

    for el in body_elements(doc):
        if el.tag == qn("w:p"):
            txt = text_of_paragraph_element(el)
            if txt:
                if (
                    re.match(r"^\d+(?:\.\d+)*\.?\s+", txt)
                    or txt in {
                        "Tổng quan",
                        "Hiện trạng",
                        "Thiết kế phần mềm",
                        "MỤC LỤC",
                        "DANH MỤC HÌNH ẢNH",
                        "DANH MỤC BẢNG",
                        "DANH MỤC VIẾT TẮT",
                    }
                ):
                    current_heading = txt
                previous_text = txt
        elif el.tag == qn("w:tbl"):
            original_table_idx += 1
            table = table_by_element[id(el)]
            if original_table_idx == 1 or original_table_idx == 24 or is_blank_table(table):
                continue
            caption_no += 1
            title = table_title(original_table_idx, table, current_heading, previous_text)
            caption = f"Bảng {caption_no}: {title}"
            insert_paragraph_before(
                el,
                doc,
                caption,
                style=doc.styles["CaptionBang"],
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
            )
            inserted.append(caption)

    return inserted


def add_figure_captions(doc):
    figure_source_idx = 0
    caption_no = 0
    inserted = []

    for el in body_elements(doc):
        if el.tag != qn("w:p"):
            continue
        blips = el.findall(".//" + qn("a:blip"))
        if not blips:
            continue
        figure_source_idx += 1
        if figure_source_idx == 1:
            continue
        title_index = figure_source_idx - 2
        if title_index >= len(FIGURE_TITLES):
            title = f"Hình minh họa {figure_source_idx - 1}"
        else:
            title = FIGURE_TITLES[title_index]
        caption_no += 1
        caption = f"Hình {caption_no}: {title}"
        insert_paragraph_after(
            el,
            doc,
            caption,
            style=doc.styles["CaptionHinh"],
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )
        inserted.append(caption)

    return inserted


def write_update_fields_setting(doc):
    settings = doc.settings._element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        settings.append(existing)
    existing.set(qn("w:val"), "true")


def main():
    if not zipfile.is_zipfile(INPUT):
        raise SystemExit(f"{INPUT} không phải file DOCX hợp lệ")

    os.makedirs("outputs", exist_ok=True)
    doc = Document(INPUT)
    ensure_caption_style(doc, "CaptionHinh")
    ensure_caption_style(doc, "CaptionBang")
    remove_generated_captions(doc)
    reset_front_matter_lists(doc)
    table_captions = add_table_captions(doc)
    figure_captions = add_figure_captions(doc)
    write_update_fields_setting(doc)
    doc.save(OUTPUT)

    print(f"saved={OUTPUT}")
    print(f"table_captions={len(table_captions)}")
    print(f"figure_captions={len(figure_captions)}")
    print(f"first_table_caption={table_captions[0] if table_captions else ''}")
    print(f"last_table_caption={table_captions[-1] if table_captions else ''}")
    print(f"first_figure_caption={figure_captions[0] if figure_captions else ''}")
    print(f"last_figure_caption={figure_captions[-1] if figure_captions else ''}")


if __name__ == "__main__":
    main()
