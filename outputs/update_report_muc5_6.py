from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "Nhom5QLTCVTHP.docx"
OUT_PATH = ROOT / "outputs" / "Nhom5QLTCVTHP_muc5_6.docx"

MEMBERS = [
    "Nguyễn Đăng Minh Quân - 24521437",
    "Phạm Nguyễn Tấn Sang - 23521346",
    "Bùi Lê Huy Phước - 23521228",
    "Đỗ Hoàng Phúc - 23521195",
]


def rows(*items):
    return list(items)


def obj(name, kind, constraint, function):
    return [name, kind, constraint, function]


def evt(event, handling):
    return [event, handling]


def business_event_label(event):
    if "API lỗi" in event:
        return event.replace("API lỗi", "Dữ liệu không tải được")
    return event


def business_flow(screen, event, handling):
    event_l = event.lower()
    screen_name = screen["name"].lower()

    if "đăng nhập" in event_l:
        return "Hệ thống kiểm tra tài khoản và mật khẩu, xác định vai trò người dùng, tạo phiên đăng nhập hợp lệ và chuyển người dùng đến khu vực chức năng tương ứng."
    if "sai tài khoản" in event_l or "không tồn tại" in event_l or "otp sai" in event_l or "lỗi" in event_l and "file" not in event_l:
        return "Hệ thống giữ nguyên dữ liệu cần thiết trên form, hiển thị lý do không hợp lệ rõ ràng để người dùng chỉnh lại và không ghi nhận thay đổi sai."
    if "hiện mật khẩu" in event_l:
        return "Giao diện đổi trạng thái ẩn/hiện mật khẩu để người dùng kiểm tra nội dung nhập nhưng không làm thay đổi dữ liệu đăng nhập."
    if "quên mật khẩu" in event_l or "gửi mã otp" in event_l:
        return "Hệ thống kiểm tra định danh tài khoản, sinh mã xác thực còn hạn sử dụng, gửi mã cho người dùng và hướng dẫn sang bước đặt lại mật khẩu."
    if "đặt lại mật khẩu" in event_l:
        return "Hệ thống kiểm tra định danh, mã OTP và hai lần nhập mật khẩu; nếu hợp lệ thì cập nhật mật khẩu mới và yêu cầu người dùng đăng nhập lại."
    if "mật khẩu xác nhận" in event_l:
        return "Giao diện so sánh hai lần nhập mật khẩu, báo lỗi ngay khi không khớp và chỉ cho tiếp tục khi dữ liệu xác nhận hợp lệ."
    if "tải" in event_l or "load" in event_l:
        return "Hệ thống xác định người dùng hoặc bộ lọc hiện tại, lấy dữ liệu liên quan, tổng hợp số liệu cần thiết và hiển thị vào đúng vùng danh sách, bảng hoặc thẻ thông tin."
    if "tìm" in event_l or "lọc" in event_l or "đổi bộ lọc" in event_l or "reset bộ lọc" in event_l:
        return "Hệ thống áp dụng tiêu chí người dùng chọn, cập nhật danh sách từ trang đầu tiên, giữ trạng thái lọc và hiển thị thông báo khi không có dữ liệu phù hợp."
    if "thêm" in event_l and "đơn" not in event_l and "môn vào chương trình" not in event_l:
        return "Người dùng nhập dữ liệu mới; hệ thống kiểm tra trường bắt buộc, định dạng, dữ liệu trùng và ràng buộc liên quan trước khi lưu và cập nhật danh sách."
    if "sửa" in event_l:
        return "Hệ thống nạp dữ liệu bản ghi vào form, khóa các trường không được phép sửa, kiểm tra thay đổi hợp lệ rồi lưu cập nhật và thông báo kết quả."
    if "xóa" in event_l or "hủy" in event_l and "hủy chốt" not in event_l:
        return "Hệ thống yêu cầu xác nhận, kiểm tra dữ liệu đang được tham chiếu hoặc trạng thái nghiệp vụ; nếu hợp lệ thì ghi nhận hủy/xóa và cập nhật lại màn hình."
    if "click dòng" in event_l or "xem chi tiết" in event_l or "xem bảng" in event_l or "xem lớp" in event_l or "xem sinh viên" in event_l:
        return "Hệ thống lấy bản ghi đang chọn, mở vùng chi tiết hoặc popup chỉ xem, trình bày đầy đủ thông tin liên quan mà không cho thay đổi dữ liệu ngoài thao tác được phép."
    if "import" in event_l or "nhập excel" in event_l:
        return "Hệ thống đọc file theo mẫu quy định, kiểm tra từng dòng, ghi nhận các dòng hợp lệ, bỏ qua dòng lỗi và trả bảng kết quả để người dùng đối chiếu."
    if "xuất excel" in event_l:
        return "Hệ thống lấy dữ liệu theo bộ lọc hiện tại, tạo file Excel đúng mẫu báo cáo và trả file cho người dùng tải xuống."
    if "upload" in event_l or "cập nhật ảnh" in event_l:
        return "Hệ thống kiểm tra định dạng và dung lượng ảnh, lưu ảnh hợp lệ, cập nhật hồ sơ và thay đổi ảnh hiển thị trên giao diện."
    if "validate lịch" in event_l or "trùng lịch" in event_l:
        return "Hệ thống so sánh thứ, tiết, phòng, giảng viên và lớp đang tồn tại để phát hiện trùng lịch trước khi cho phép lưu lịch học."
    if "mở lớp" in event_l:
        return "Hệ thống kiểm tra học kỳ, môn học, lịch học, giảng viên, phòng và sĩ số; nếu hợp lệ thì chuyển lớp sang trạng thái được mở cho sinh viên đăng ký."
    if "đóng lớp" in event_l:
        return "Hệ thống kiểm tra lớp mở và dữ liệu đăng ký liên quan, sau đó cập nhật trạng thái đóng lớp và không cho phát sinh đăng ký mới."
    if "chọn môn" in event_l:
        return "Người dùng tra cứu môn trong popup, chọn đúng bản ghi; hệ thống đưa mã và tên môn về form nghiệp vụ để tránh chọn nhầm từ danh sách dài."
    if "thêm môn vào chương trình" in event_l:
        return "Hệ thống kiểm tra ngành, môn học và học kỳ dự kiến; nếu môn chưa tồn tại trong chương trình thì ghi nhận vào cấu trúc đào tạo của ngành."
    if "đăng ký" in event_l and "chốt" not in event_l:
        return "Hệ thống kiểm tra cửa sổ đăng ký, sĩ số lớp, môn trùng, lịch trùng, điều kiện tiên quyết và giới hạn tín chỉ; hợp lệ thì ghi vào phiếu đăng ký và cập nhật học phí tạm tính."
    if "chốt đăng ký" in event_l:
        return "Hệ thống kiểm tra thời hạn đăng ký và đơn cứu xét còn chờ; nếu đủ điều kiện thì khóa dữ liệu đăng ký của học kỳ để chuyển sang bước thu học phí."
    if "hủy chốt đăng ký" in event_l:
        return "Hệ thống kiểm tra trạng thái học kỳ và quyền thao tác, sau đó mở lại dữ liệu đăng ký để quản trị viên tiếp tục xử lý điều chỉnh."
    if "duyệt đơn" in event_l:
        return "Hệ thống kiểm tra thời gian cứu xét, lý do, lớp cần thêm hoặc hủy, sĩ số và lịch học; hợp lệ thì cập nhật phiếu đăng ký và trạng thái đơn."
    if "từ chối đơn" in event_l:
        return "Quản trị viên nhập lý do từ chối; hệ thống lưu kết quả xử lý, giữ nguyên phiếu đăng ký và thông báo trạng thái cho sinh viên."
    if "gửi đơn" in event_l:
        return "Sinh viên nhập lý do cứu xét; hệ thống kiểm tra thời gian tiếp nhận, dữ liệu lớp liên quan và tạo đơn ở trạng thái chờ duyệt."
    if "thanh toán" in event_l or "tạo thanh toán" in event_l:
        return "Hệ thống kiểm tra phiếu thu còn hiệu lực, số tiền còn nợ và phương thức thanh toán; sau đó tạo giao dịch, ghi nhận trạng thái và cập nhật công nợ."
    if "xác nhận" in event_l or "thất bại" in event_l or "hủy phiếu" in event_l:
        return "Hệ thống cập nhật trạng thái phiếu thu theo kết quả xử lý, ghi nhận thời điểm xác nhận và tính lại số tiền đã đóng, còn nợ của sinh viên."
    if "gán quyền" in event_l or "kiểm tra truy cập" in event_l:
        return "Hệ thống cập nhật quyền cho nhóm người dùng, áp dụng lại vào sidebar và kiểm tra quyền trước khi cho mở từng màn hình quản trị."
    if "khôi phục" in event_l or "xóa vĩnh viễn" in event_l:
        return "Hệ thống kiểm tra bản ghi trong thùng rác, xác nhận thao tác, khôi phục dữ liệu về danh mục hoặc xóa vĩnh viễn theo quyền hiện có."
    if "không có dữ liệu" in event_l or "dữ liệu rỗng" in event_l:
        return "Giao diện giữ cấu trúc bảng hoặc thẻ thống kê, hiển thị trạng thái trống rõ ràng để người dùng biết không có bản ghi phù hợp."
    if "render biểu đồ" in event_l or "cập nhật biểu đồ" in event_l:
        return "Hệ thống tổng hợp số liệu theo kỳ hoặc theo tháng, chuyển thành dữ liệu biểu đồ và làm mới phần trực quan hóa trên màn hình."
    if "in báo cáo" in event_l:
        return "Hệ thống dùng dữ liệu báo cáo đang hiển thị để mở chế độ in, giữ tiêu đề, bộ lọc và bảng số liệu phục vụ lưu trữ."
    if "tham số" in screen_name or "giá trị số" in event_l:
        return "Hệ thống kiểm tra giá trị tham số, chuẩn hóa dữ liệu nhập, lưu cấu hình mới và áp dụng cho các bước nghiệp vụ bị ảnh hưởng."
    return "Hệ thống tiếp nhận thao tác, kiểm tra điều kiện nghiệp vụ, cập nhật dữ liệu liên quan nếu hợp lệ và phản hồi kết quả ngay trên màn hình."


screens = [
    {
        "name": "Màn hình đăng nhập",
        "type": "hệ thống",
        "member": MEMBERS[0],
        "route": "/login, /admin/login",
        "view": "src/views/pages/login.pug",
        "js": "src/public/js/login.js",
        "function": "Xác thực tài khoản sinh viên hoặc quản trị viên, lưu token và điều hướng về dashboard đúng vai trò.",
        "layout": "Màn hình auth có logo, tiêu đề theo vai trò đăng nhập, form tên đăng nhập, mật khẩu, nút hiện mật khẩu, nút đăng nhập và liên kết quên mật khẩu.",
        "objects": rows(
            obj("Tiêu đề đăng nhập", "vùng hiển thị", "Hiển thị theo loginRole sinh viên hoặc admin", "Cho người dùng biết đang đăng nhập vào cổng sinh viên hay khu vực quản trị."),
            obj("Tên đăng nhập", "input text", "Bắt buộc nhập, autocomplete username", "Nhập tài khoản để gửi đến API đăng nhập."),
            obj("Mật khẩu", "input password", "Bắt buộc nhập, có nút ẩn hiện mật khẩu", "Nhập mật khẩu và hỗ trợ kiểm tra lại ký tự đã nhập."),
            obj("Nút Đăng nhập", "button submit", "Chỉ xử lý khi form hợp lệ", "Gửi POST đến /api/auth/login hoặc /api/auth/admin/login.")
        ),
        "events": rows(
            evt("Nhấn Đăng nhập", "Gửi thông tin đăng nhập, kiểm tra role trả về, lưu token cookie và chuyển đến /student/dashboard hoặc /admin/dashboard."),
            evt("Nhập sai tài khoản hoặc mật khẩu", "Hiển thị thông báo lỗi từ API, khôi phục trạng thái nút đăng nhập."),
            evt("Nhấn nút hiện mật khẩu", "Đổi type của ô mật khẩu giữa password và text, cập nhật biểu tượng và aria-label."),
            evt("Chọn Quên mật khẩu", "Điều hướng sang màn hình quên mật khẩu tương ứng với vai trò.")
        ),
        "install": "Cài view login.pug, login.js, API /api/auth/login và /api/auth/admin/login, kiểm tra redirect theo token.",
        "tests": "Kiểm thử đăng nhập sinh viên, đăng nhập admin, sai mật khẩu, sai vai trò và nút hiện mật khẩu.",
        "completion": "95%",
        "proof": "[Chụp đăng nhập sinh viên thành công], [Chụp đăng nhập admin thành công], [Chụp lỗi sai mật khẩu]"
    },
    {
        "name": "Màn hình quên mật khẩu",
        "type": "hệ thống",
        "member": MEMBERS[0],
        "route": "/forgot-password, /admin/forgot-password",
        "view": "src/views/pages/forgot-password.pug",
        "js": "src/public/js/login.js",
        "function": "Nhận tên đăng nhập hoặc email, gửi yêu cầu cấp OTP đặt lại mật khẩu.",
        "layout": "Màn hình auth gồm tiêu đề quên mật khẩu, ô tên đăng nhập hoặc email, nút gửi mã OTP và vùng thông báo kết quả.",
        "objects": rows(
            obj("Tên đăng nhập hoặc email", "input text", "Bắt buộc nhập", "Xác định tài khoản cần cấp OTP."),
            obj("Nút Gửi mã OTP", "button submit", "Chỉ gửi khi có định danh", "Gọi API /api/auth/forgot-password."),
            obj("Vùng thông báo", "message area", "Hiển thị thành công hoặc lỗi", "Thông báo trạng thái gửi OTP cho người dùng."),
            obj("Liên kết quay lại đăng nhập", "link", "Theo vai trò hiện tại", "Quay về /login hoặc /admin/login.")
        ),
        "events": rows(
            evt("Nhấn Gửi mã OTP", "Gửi định danh và role đến API quên mật khẩu, sau đó chuyển sang màn hình đặt lại mật khẩu nếu thành công."),
            evt("Định danh không tồn tại", "Hiển thị lỗi từ API và giữ nguyên dữ liệu đã nhập."),
            evt("Để trống ô định danh", "Trình duyệt chặn submit vì trường required."),
            evt("Quay lại đăng nhập", "Điều hướng về màn hình đăng nhập tương ứng.")
        ),
        "install": "Cài forgot-password.pug, handler handleForgotPassword và API /api/auth/forgot-password dùng Redis OTP.",
        "tests": "Kiểm thử gửi OTP bằng username, gửi OTP bằng email, định danh không tồn tại và trường bắt buộc.",
        "completion": "95%",
        "proof": "[Chụp gửi OTP thành công], [Chụp lỗi tài khoản không tồn tại], [Chụp chuyển sang đặt lại mật khẩu]"
    },
    {
        "name": "Màn hình đặt lại mật khẩu",
        "type": "hệ thống",
        "member": MEMBERS[0],
        "route": "/reset-password",
        "view": "src/views/pages/reset-password.pug",
        "js": "src/public/js/login.js",
        "function": "Xác nhận OTP và cập nhật mật khẩu mới cho tài khoản sinh viên hoặc quản trị viên.",
        "layout": "Màn hình auth có hidden role, ô định danh, OTP, mật khẩu mới, xác nhận mật khẩu mới, nút đặt lại mật khẩu.",
        "objects": rows(
            obj("Role đặt lại", "hidden input", "Lấy từ query hoặc luồng quên mật khẩu", "Xác định API xử lý theo vai trò."),
            obj("Mã OTP", "input text", "Bắt buộc nhập, tối đa 8 ký tự", "Nhập mã xác nhận nhận được từ email."),
            obj("Mật khẩu mới", "input password", "Bắt buộc, minlength 6", "Nhập mật khẩu mới."),
            obj("Nhập lại mật khẩu mới", "input password", "Bắt buộc, phải khớp mật khẩu mới", "Kiểm tra xác nhận trước khi gửi API.")
        ),
        "events": rows(
            evt("Nhấn Đặt lại mật khẩu", "Kiểm tra hai mật khẩu khớp nhau rồi gọi /api/auth/reset-password."),
            evt("OTP sai hoặc hết hạn", "Hiển thị lỗi, không xóa dữ liệu định danh."),
            evt("Mật khẩu xác nhận không khớp", "Chặn gửi API và hiển thị thông báo lỗi tại form."),
            evt("Đặt lại thành công", "Reset form và hướng người dùng quay lại màn hình đăng nhập.")
        ),
        "install": "Cài reset-password.pug, handler handleResetPassword và API /api/auth/reset-password kiểm tra OTP Redis.",
        "tests": "Kiểm thử OTP đúng, OTP sai, mật khẩu ngắn, xác nhận mật khẩu không khớp và reset thành công.",
        "completion": "95%",
        "proof": "[Chụp đặt lại mật khẩu thành công], [Chụp lỗi OTP], [Chụp lỗi xác nhận mật khẩu]"
    },
    {
        "name": "Màn hình bảng điều khiển quản trị",
        "type": "báo cáo",
        "member": MEMBERS[0],
        "route": "/admin/dashboard",
        "view": "src/views/pages/admin/dashboard.pug",
        "js": "src/public/js/admin/dashboard.js",
        "function": "Tổng hợp số liệu sinh viên, môn học, doanh thu, đăng ký, sinh viên còn nợ và hoạt động gần đây.",
        "layout": "Màn hình admin gồm các thẻ chỉ số, biểu đồ doanh thu theo tháng, biểu đồ đăng ký theo học kỳ, bảng sinh viên nợ học phí và danh sách hoạt động gần đây.",
        "objects": rows(
            obj("Thẻ chỉ số tổng quan", "stat card", "Dữ liệu lấy từ /api/dashboard/stats", "Hiển thị nhanh số sinh viên, môn học, đăng ký và công nợ."),
            obj("Biểu đồ doanh thu", "chart", "Dùng Chart.js", "Trực quan hóa doanh thu từng tháng."),
            obj("Biểu đồ đăng ký", "chart", "Dữ liệu theo học kỳ", "Theo dõi số lượt đăng ký học phần."),
            obj("Bảng sinh viên nợ học phí", "table", "Cột MSSV, họ tên, học kỳ, còn nợ", "Liệt kê các sinh viên còn công nợ nổi bật.")
        ),
        "events": rows(
            evt("Tải trang dashboard", "Gọi /api/dashboard/stats, /api/dashboard/revenue-monthly, /api/dashboard/registration-by-semester, /api/dashboard/students-owing và /api/dashboard/recent-activity."),
            evt("API thống kê lỗi", "Giữ khung giao diện và hiển thị giá trị mặc định để không làm hỏng trang."),
            evt("Cập nhật biểu đồ", "Đổ dữ liệu vào Chart.js, tạo biểu đồ doanh thu và đăng ký."),
            evt("Xem bảng nợ", "Render danh sách sinh viên còn nợ để admin chuyển sang công nợ hoặc báo cáo chi tiết.")
        ),
        "install": "Cài dashboard.pug, dashboard.js, dashboardRoutes và dashboardController cho thống kê tổng quan.",
        "tests": "Kiểm thử tải dashboard, hiển thị thẻ số liệu, hiển thị biểu đồ, bảng sinh viên nợ và xử lý khi API rỗng.",
        "completion": "95%",
        "proof": "[Chụp dashboard admin], [Chụp biểu đồ doanh thu], [Chụp bảng sinh viên nợ]"
    },
    {
        "name": "Màn hình quản lý tỉnh/thành phố",
        "type": "danh mục",
        "member": MEMBERS[0],
        "route": "/admin/locations/provinces",
        "view": "src/views/pages/admin/locations-provinces.pug",
        "js": "src/public/js/admin/locations.js",
        "function": "Quản lý danh mục tỉnh/thành phố, tìm kiếm theo mã, tên, loại và lọc trạng thái.",
        "layout": "Màn hình admin có toolbar tìm kiếm, lọc loại và trạng thái, bảng tỉnh/thành phố, modal thêm sửa và modal chi tiết.",
        "objects": rows(
            obj("Tìm theo và từ khóa", "select + input", "Tiêu chí gồm mã, tên, loại hoặc tất cả", "Lọc danh sách tỉnh/thành phố."),
            obj("Bộ lọc loại, trạng thái", "select", "Loại tỉnh hoặc thành phố, trạng thái hoạt động", "Thu hẹp dữ liệu trên bảng."),
            obj("Bảng tỉnh/thành phố", "table", "Cột mã, tên, loại, số phường xã, trạng thái, sửa bởi, sửa lúc", "Hiển thị dữ liệu danh mục."),
            obj("Modal tỉnh/thành phố", "modal form", "Mã tỉnh là khóa chính, không sửa sau khi tạo", "Thêm hoặc cập nhật thông tin tỉnh/thành phố.")
        ),
        "events": rows(
            evt("Tìm kiếm hoặc lọc", "Cập nhật query và render lại danh sách theo /admin/locations/provinces."),
            evt("Thêm tỉnh/thành phố", "Mở modal add, gửi POST /api/locations/provinces."),
            evt("Sửa tỉnh/thành phố", "Khóa mã tỉnh, gửi PUT /api/locations/provinces/:id."),
            evt("Xóa tỉnh/thành phố", "Gọi DELETE /api/locations/provinces/:id và kiểm tra dữ liệu phường/xã tham chiếu.")
        ),
        "install": "Cài view locations-provinces.pug, locationRoutes, locationController và locations.js cho CRUD tỉnh.",
        "tests": "Kiểm thử thêm, sửa, xóa, tìm theo loại, lọc trạng thái, click dòng xem chi tiết và lỗi bản ghi đang được tham chiếu.",
        "completion": "95%",
        "proof": "[Chụp thêm tỉnh thành công], [Chụp lỗi xóa tỉnh đang có phường xã], [Chụp popup chi tiết tỉnh]"
    },
    {
        "name": "Màn hình quản lý phường/xã",
        "type": "danh mục",
        "member": MEMBERS[0],
        "route": "/admin/locations/wards",
        "view": "src/views/pages/admin/locations-wards.pug",
        "js": "src/public/js/admin/locations.js",
        "function": "Quản lý danh mục phường/xã theo tỉnh, loại, khu vực và trạng thái.",
        "layout": "Màn hình admin gồm vùng tìm kiếm, lọc tỉnh, lọc loại phường/xã, lọc khu vực, bảng danh sách, modal thêm sửa và modal chi tiết.",
        "objects": rows(
            obj("Bộ lọc địa bàn", "select group", "Tỉnh, loại, khu vực và trạng thái", "Lọc danh sách phường/xã."),
            obj("Bảng phường/xã", "table", "Cột mã, tên, loại, tỉnh, khu vực, trạng thái, sửa bởi, sửa lúc", "Hiển thị danh mục địa chỉ cấp phường/xã."),
            obj("Modal phường/xã", "modal form", "Mã phường/xã là khóa chính, tỉnh bắt buộc", "Thêm hoặc cập nhật thông tin phường/xã."),
            obj("Nút xóa", "button", "Có xác nhận trước khi xóa", "Xóa mềm phường/xã nếu không vi phạm ràng buộc.")
        ),
        "events": rows(
            evt("Lọc theo tỉnh hoặc khu vực", "Cập nhật query và tải lại danh sách phường/xã phù hợp."),
            evt("Thêm phường/xã", "Gửi POST /api/locations/wards với mã, tên, tỉnh, loại, khu vực."),
            evt("Sửa phường/xã", "Gửi PUT /api/locations/wards/:id, mã không được sửa."),
            evt("Click dòng dữ liệu", "Mở modal chi tiết chỉ xem thông tin phường/xã.")
        ),
        "install": "Cài view locations-wards.pug, API /api/locations/wards và xử lý lọc trong locations.js.",
        "tests": "Kiểm thử thêm, sửa, xóa, tìm theo loại, lọc tỉnh, lọc khu vực, xem chi tiết và lỗi dữ liệu bắt buộc.",
        "completion": "95%",
        "proof": "[Chụp thêm phường xã thành công], [Chụp lọc theo tỉnh], [Chụp lỗi thiếu tỉnh]"
    },
    {
        "name": "Màn hình quản lý năm học",
        "type": "danh mục",
        "member": MEMBERS[0],
        "route": "/admin/academic-years",
        "view": "src/views/pages/admin/academic-years.pug",
        "js": "src/public/js/admin/academic-years.js",
        "function": "Quản lý năm học, tự đồng bộ năm bắt đầu và năm kết thúc theo mã năm học.",
        "layout": "Màn hình có toolbar tìm kiếm, lọc trạng thái, bảng năm học và modal thêm sửa gồm mã, tên, năm bắt đầu, năm kết thúc, trạng thái.",
        "objects": rows(
            obj("Ô tìm kiếm năm học", "select + input", "Tìm theo mã hoặc tên năm học", "Lọc bảng năm học."),
            obj("Bảng năm học", "table", "Cột mã, tên, năm bắt đầu, năm kết thúc, trạng thái, sửa bởi, sửa lúc", "Hiển thị danh sách năm học."),
            obj("Form năm học", "modal form", "Mã năm học không sửa sau khi tạo", "Nhập thông tin năm học."),
            obj("Nút Lưu", "button", "Kiểm tra năm kết thúc lớn hơn năm bắt đầu", "Tạo hoặc cập nhật năm học.")
        ),
        "events": rows(
            evt("Nhập mã năm học", "Hàm syncAcademicYearFromCode tự suy ra năm bắt đầu và năm kết thúc nếu mã đúng định dạng."),
            evt("Thêm năm học", "Gửi POST /api/semesters/years."),
            evt("Sửa năm học", "Gửi PUT /api/semesters/years/:id và khóa mã năm học."),
            evt("Xóa năm học", "Gọi DELETE /api/semesters/years/:id và kiểm tra học kỳ đang tham chiếu.")
        ),
        "install": "Cài academic-years.pug, academic-years.js và các endpoint /api/semesters/years.",
        "tests": "Kiểm thử thêm, sửa, xóa năm học, tự đồng bộ năm, tìm kiếm, lọc trạng thái và lỗi năm không hợp lệ.",
        "completion": "95%",
        "proof": "[Chụp thêm năm học thành công], [Chụp lỗi năm kết thúc], [Chụp xóa năm học]"
    },
    {
        "name": "Màn hình quản lý học kỳ",
        "type": "danh mục",
        "member": MEMBERS[0],
        "route": "/admin/semesters",
        "view": "src/views/pages/admin/semesters.pug",
        "js": "src/public/js/admin/semesters.js",
        "function": "Quản lý học kỳ, mốc đăng ký, mốc cứu xét và mốc thu học phí.",
        "layout": "Màn hình có bộ lọc học kỳ, trạng thái, chốt đăng ký, mở thu, mốc thời gian; bảng học kỳ và modal nhập đầy đủ ngày bắt đầu kết thúc.",
        "objects": rows(
            obj("Bộ lọc học kỳ", "filter controls", "Lọc theo loại học kỳ, trạng thái, chốt đăng ký, mở thu", "Tìm nhanh học kỳ cần quản lý."),
            obj("Bảng học kỳ", "table", "Hiển thị năm học, loại học kỳ, trạng thái và các mốc nghiệp vụ", "Theo dõi tiến độ học kỳ."),
            obj("Form học kỳ", "modal form", "Mã học kỳ bắt buộc, ngày kết thúc sau ngày bắt đầu", "Thêm hoặc sửa học kỳ."),
            obj("Các mốc nghiệp vụ", "date inputs", "Đăng ký, cứu xét, thu học phí có thứ tự thời gian hợp lệ", "Điều khiển luồng đăng ký và thanh toán.")
        ),
        "events": rows(
            evt("Lọc học kỳ", "Gọi loadSemesters với query lọc và render lại bảng."),
            evt("Thêm hoặc sửa học kỳ", "Validate các cặp ngày rồi gửi POST hoặc PUT /api/semesters."),
            evt("Reset bộ lọc", "Xóa query lọc và tải lại danh sách mặc định."),
            evt("Click dòng học kỳ", "Hiển thị chi tiết học kỳ để xem các mốc thời gian.")
        ),
        "install": "Cài semesters.pug, semesters.js, semesterRoutes và semesterController cho quản lý học kỳ.",
        "tests": "Kiểm thử thêm, sửa, xóa học kỳ, lọc trạng thái, lọc chốt đăng ký, lọc mở thu và lỗi mốc thời gian không hợp lệ.",
        "completion": "95%",
        "proof": "[Chụp thêm học kỳ thành công], [Chụp lỗi ngày học kỳ], [Chụp lọc trạng thái mở thu]"
    },
    {
        "name": "Màn hình quản lý tiết học",
        "type": "danh mục",
        "member": MEMBERS[0],
        "route": "/admin/periods",
        "view": "src/views/pages/admin/periods.pug",
        "js": "src/public/js/admin/periods.js",
        "function": "Quản lý tiết học gồm mã tiết, tên tiết, thứ tự, giờ bắt đầu, giờ kết thúc và trạng thái.",
        "layout": "Màn hình có tìm kiếm theo mã hoặc tên tiết, bảng danh sách tiết học và modal thêm sửa.",
        "objects": rows(
            obj("Tìm theo tiết", "select + input", "Tìm theo mã tiết hoặc tên tiết", "Lọc danh sách tiết học."),
            obj("Bảng tiết học", "table", "Cột mã tiết, tên tiết, thứ tự, bắt đầu, kết thúc, trạng thái", "Hiển thị khung thời gian học."),
            obj("Form tiết học", "modal form", "Mã tiết, tên tiết, thứ tự, giờ bắt đầu và giờ kết thúc bắt buộc", "Thêm hoặc cập nhật tiết học."),
            obj("Trạng thái tiết", "select", "Đang dùng hoặc tạm khóa", "Ẩn hiện tiết trong các màn hình chọn lịch.")
        ),
        "events": rows(
            evt("Thêm tiết học", "Gửi POST /api/periods sau khi kiểm tra dữ liệu bắt buộc."),
            evt("Sửa tiết học", "Gửi PUT /api/periods/:id và khóa mã tiết khi cần."),
            evt("Xóa tiết học", "Gọi DELETE /api/periods/:id và kiểm tra lớp học đang dùng tiết."),
            evt("Tìm kiếm", "Cập nhật query tìm kiếm và render danh sách tiết học.")
        ),
        "install": "Cài periods.pug, periods.js, periodRoutes và periodController cho CRUD tiết học.",
        "tests": "Kiểm thử thêm, sửa, xóa tiết học, tìm theo mã, tìm theo tên, lỗi giờ kết thúc trước giờ bắt đầu và ràng buộc lớp học.",
        "completion": "95%",
        "proof": "[Chụp thêm tiết học thành công], [Chụp lỗi giờ tiết học], [Chụp tìm kiếm tiết học]"
    },
    {
        "name": "Màn hình quản lý khoa",
        "type": "danh mục",
        "member": MEMBERS[0],
        "route": "/admin/faculties",
        "view": "src/views/pages/admin/faculties.pug",
        "js": "src/public/js/admin/faculties.js",
        "function": "Quản lý khoa, thông tin liên hệ, trưởng khoa và số ngành, số môn liên quan.",
        "layout": "Màn hình có tìm kiếm theo mã, tên, viết tắt; bảng khoa; modal thêm sửa gồm mã, tên, viết tắt, trưởng khoa, email, số điện thoại, địa chỉ.",
        "objects": rows(
            obj("Tìm kiếm khoa", "select + input", "Tìm theo mã khoa, tên khoa hoặc tên viết tắt", "Lọc danh sách khoa."),
            obj("Bảng khoa", "table", "Cột mã, tên, viết tắt, trưởng khoa, liên hệ, số ngành, số môn", "Hiển thị thông tin khoa."),
            obj("Form khoa", "modal form", "Mã khoa và tên khoa bắt buộc, email đúng định dạng", "Thêm hoặc cập nhật khoa."),
            obj("Nút xóa khoa", "button", "Có xác nhận", "Xóa mềm khoa nếu không bị ngành, môn học hoặc giảng viên tham chiếu.")
        ),
        "events": rows(
            evt("Thêm khoa", "Gửi POST /api/faculties với dữ liệu form."),
            evt("Sửa khoa", "Gửi PUT /api/faculties/:id, mã khoa không đổi."),
            evt("Xóa khoa", "Gọi DELETE /api/faculties/:id và xử lý lỗi ràng buộc."),
            evt("Click dòng", "Hiển thị chi tiết khoa theo dữ liệu bản ghi.")
        ),
        "install": "Cài faculties.pug, faculties.js, facultyRoutes và facultyController.",
        "tests": "Kiểm thử thêm, sửa, xóa khoa, tìm theo viết tắt, lỗi email và lỗi khoa đang được tham chiếu.",
        "completion": "95%",
        "proof": "[Chụp thêm khoa thành công], [Chụp lỗi email khoa], [Chụp chi tiết khoa]"
    },
    {
        "name": "Màn hình quản lý ngành học",
        "type": "danh mục",
        "member": MEMBERS[0],
        "route": "/admin/majors",
        "view": "src/views/pages/admin/majors.pug",
        "js": "src/public/js/admin/majors.js",
        "function": "Quản lý ngành học và xem chương trình đào tạo theo ngành.",
        "layout": "Màn hình có tìm kiếm theo mã ngành, tên ngành, khoa; bảng ngành học; modal ngành; khu vực chương trình đào tạo và modal thêm môn vào chương trình.",
        "objects": rows(
            obj("Bộ lọc ngành", "select + input", "Tìm theo mã ngành, tên ngành hoặc tên khoa", "Lọc danh sách ngành học."),
            obj("Bảng ngành học", "table", "Cột mã ngành, tên ngành, khoa, số tín chỉ tối thiểu, thời gian đào tạo", "Hiển thị dữ liệu ngành."),
            obj("Modal ngành học", "modal form", "Mã ngành và khoa bắt buộc", "Thêm hoặc cập nhật ngành học."),
            obj("Khu vực CTĐT", "table + modal", "Chọn ngành trước khi thêm môn", "Xem và cập nhật chương trình đào tạo của ngành.")
        ),
        "events": rows(
            evt("Thêm ngành học", "Gửi POST /api/majors với thông tin ngành."),
            evt("Sửa ngành học", "Gửi PUT /api/majors/:id và khóa mã ngành."),
            evt("Xem chương trình đào tạo", "Gọi /api/majors/:id/curriculum hoặc /api/majors/curriculum/items để render danh sách môn."),
            evt("Thêm môn vào CTĐT", "Chọn môn, học kỳ dự kiến, trạng thái và lưu vào chương trình.")
        ),
        "install": "Cài majors.pug, majors.js, majorRoutes và majorController cho ngành học và CTĐT.",
        "tests": "Kiểm thử thêm, sửa, xóa ngành, tìm theo khoa, mở CTĐT, thêm môn vào CTĐT và lỗi trùng mã ngành.",
        "completion": "95%",
        "proof": "[Chụp thêm ngành thành công], [Chụp CTĐT theo ngành], [Chụp lỗi trùng ngành]"
    },
    {
        "name": "Màn hình quản lý sinh viên",
        "type": "danh mục",
        "member": MEMBERS[1],
        "route": "/admin/students",
        "view": "src/views/pages/admin/students.pug",
        "js": "src/public/js/admin/students.js",
        "function": "Quản lý hồ sơ sinh viên, ảnh đại diện, ngành, địa chỉ, dân tộc, đối tượng ưu tiên và import export Excel.",
        "layout": "Màn hình admin có tìm kiếm, lọc ngành, trạng thái, bảng sinh viên, modal thêm sửa nhiều trường, chọn ảnh và nhập xuất Excel.",
        "objects": rows(
            obj("Tìm kiếm sinh viên", "select + input", "Tìm theo MSSV, họ tên hoặc email", "Lọc danh sách sinh viên."),
            obj("Bảng sinh viên", "table", "Cột MSSV, họ tên, liên hệ, ngành, khoa, dân tộc, đối tượng, trạng thái", "Hiển thị hồ sơ sinh viên."),
            obj("Form sinh viên", "modal form", "MSSV, họ tên, ngày sinh, CCCD, dân tộc, địa chỉ, ngành bắt buộc", "Thêm hoặc sửa hồ sơ sinh viên."),
            obj("Import, export và ảnh", "file controls", "File Excel hoặc file ảnh đúng định dạng", "Nhập danh sách, xuất Excel và cập nhật avatar sinh viên.")
        ),
        "events": rows(
            evt("Thêm sinh viên", "Gửi POST /api/students, kiểm tra email, số điện thoại, dân tộc và đối tượng ưu tiên."),
            evt("Sửa sinh viên", "Gửi PUT /api/students/:id, cập nhật các thông tin hồ sơ."),
            evt("Nhập Excel", "Gửi file đến /api/students/import và hiển thị kết quả từng dòng."),
            evt("Cập nhật ảnh", "Gửi file ảnh đến /api/students/:id/avatar.")
        ),
        "install": "Cài students.pug, students.js, studentRoutes, studentController và avatarUpload middleware.",
        "tests": "Kiểm thử thêm, sửa, xóa, import Excel, export Excel, upload ảnh, tìm kiếm và lỗi trùng MSSV hoặc CCCD.",
        "completion": "95%",
        "proof": "[Chụp thêm sinh viên thành công], [Chụp import Excel sinh viên], [Chụp lỗi trùng MSSV]"
    },
    {
        "name": "Màn hình quản lý môn học",
        "type": "danh mục",
        "member": MEMBERS[1],
        "route": "/admin/courses",
        "view": "src/views/pages/admin/courses.pug",
        "js": "src/public/js/admin/courses.js",
        "function": "Quản lý danh mục môn học, số tiết, loại môn, số tín chỉ tự tính, import và export Excel.",
        "layout": "Màn hình có tìm kiếm theo mã, tên, loại, khoa; lọc khoa và loại môn; bảng môn học; modal thêm sửa; modal chi tiết; modal nhập Excel.",
        "objects": rows(
            obj("Bộ lọc môn học", "select + input", "Tìm theo mã, tên, loại môn hoặc khoa", "Lọc danh mục môn học."),
            obj("Bảng môn học", "table", "Cột mã môn, tên, tín chỉ, loại, số tiết, khoa, sửa bởi, sửa lúc", "Hiển thị danh sách môn học."),
            obj("Form môn học", "modal form", "Mã, tên, khoa, loại môn, số tiết bắt buộc", "Thêm hoặc sửa môn học."),
            obj("Số tín chỉ", "readonly input", "Không sửa trực tiếp, tính theo số tiết và loại môn", "Giảm sai lệch khi nhập dữ liệu môn học.")
        ),
        "events": rows(
            evt("Thêm hoặc sửa môn học", "Gọi POST hoặc PUT /api/courses, tự đồng bộ tín chỉ bằng syncCredits."),
            evt("Xóa môn học", "Gọi DELETE /api/courses/:id và kiểm tra ràng buộc lớp, CTĐT, môn mở."),
            evt("Nhập Excel môn học", "Gửi file đến /api/courses/import và hiển thị kết quả thành công hoặc lỗi."),
            evt("Xuất Excel", "Gọi /api/courses/export với bộ lọc hiện tại.")
        ),
        "install": "Cài courses.pug, courses.js, courseRoutes và courseController cho CRUD và import export.",
        "tests": "Kiểm thử thêm, sửa, xóa môn học, tự tính tín chỉ, import thành công, import lỗi từng dòng và export Excel.",
        "completion": "95%",
        "proof": "[Chụp thêm môn học thành công], [Chụp số tín chỉ tự tính], [Chụp lỗi import môn học]"
    },
    {
        "name": "Màn hình quản lý môn học mở",
        "type": "nghiệp vụ",
        "member": MEMBERS[1],
        "route": "/admin/open-courses",
        "view": "src/views/pages/admin/open-courses.pug",
        "js": "src/public/js/admin/open-courses.js",
        "function": "Mở môn học theo học kỳ, chọn môn bằng popup, lọc theo học kỳ, khoa và trạng thái.",
        "layout": "Màn hình có toolbar lọc học kỳ, khoa, trạng thái, bảng môn học mở, modal thêm sửa và popup chọn môn học.",
        "objects": rows(
            obj("Bộ lọc môn mở", "select + input", "Lọc theo học kỳ, khoa, trạng thái và tiêu chí tìm kiếm", "Tìm môn học mở theo nhu cầu."),
            obj("Bảng môn học mở", "table", "Cột học kỳ, mã môn, tên môn, loại môn, khoa, trạng thái, cập nhật", "Hiển thị môn đã mở."),
            obj("Popup chọn môn học", "modal picker", "Tìm theo mã hoặc tên môn", "Chọn môn thay cho select dài."),
            obj("Form môn mở", "modal form", "Học kỳ và môn học bắt buộc, không trùng cặp học kỳ môn", "Tạo hoặc cập nhật môn học mở.")
        ),
        "events": rows(
            evt("Chọn môn từ popup", "Gọi /api/open-courses/available và gán mã môn vào form."),
            evt("Thêm môn học mở", "Gửi POST /api/open-courses với học kỳ và mã môn."),
            evt("Sửa môn học mở", "Gửi PUT /api/open-courses/:id để cập nhật trạng thái hoặc ghi chú."),
            evt("Xóa môn học mở", "Gọi DELETE /api/open-courses/:id và kiểm tra lớp mở hoặc đăng ký liên quan.")
        ),
        "install": "Cài open-courses.pug, open-courses.js, openCourseRoutes và openCourseController.",
        "tests": "Kiểm thử thêm, sửa, xóa môn mở, lọc học kỳ, popup chọn môn, lỗi trùng môn trong học kỳ và ràng buộc đăng ký.",
        "completion": "95%",
        "proof": "[Chụp thêm môn mở thành công], [Chụp popup chọn môn], [Chụp lỗi trùng môn mở]"
    },
    {
        "name": "Màn hình quản lý lớp học",
        "type": "nghiệp vụ",
        "member": MEMBERS[1],
        "route": "/admin/classes",
        "view": "src/views/pages/admin/classes.pug",
        "js": "src/public/js/admin/classes.js",
        "function": "Quản lý lớp học, lịch học, giảng viên, phòng học, sĩ số và trạng thái mở lớp.",
        "layout": "Màn hình có bộ lọc học kỳ, trạng thái mở, sắp xếp sĩ số, bảng lớp, modal lớp, popup chọn môn, giảng viên, phòng và modal danh sách sinh viên.",
        "objects": rows(
            obj("Bộ lọc lớp học", "filter controls", "Lọc theo học kỳ, trạng thái mở và sĩ số", "Tìm lớp học theo điều kiện quản lý."),
            obj("Bảng lớp học", "table", "Cột lớp, môn học, phụ trách lịch, sĩ số, trạng thái mở, cập nhật", "Hiển thị danh sách lớp."),
            obj("Form lớp học", "modal form", "Mã lớp, tên lớp, môn học bắt buộc", "Thêm hoặc sửa thông tin lớp."),
            obj("Lịch học", "select group", "Thứ, tiết bắt đầu, tiết kết thúc, phòng hợp lệ", "Kiểm tra trùng lịch và xếp phòng.")
        ),
        "events": rows(
            evt("Validate lịch học", "Gọi /api/classes/validate-schedule để kiểm tra trùng lịch giảng viên, phòng và tiết."),
            evt("Mở lớp", "Gửi POST /api/classes/open để mở lớp trong học kỳ."),
            evt("Đóng lớp", "Gọi DELETE /api/classes/opened/:id để đóng lớp đã mở."),
            evt("Xem sinh viên trong lớp", "Gọi /api/classes/:id/students theo học kỳ và hiển thị modal danh sách.")
        ),
        "install": "Cài classes.pug, classes.js, classRoutes và classController cho lớp, lịch, mở lớp.",
        "tests": "Kiểm thử thêm, sửa, xóa lớp, mở lớp, đóng lớp, xem danh sách sinh viên, sắp xếp sĩ số và lỗi trùng lịch.",
        "completion": "95%",
        "proof": "[Chụp mở lớp thành công], [Chụp lỗi trùng lịch], [Chụp danh sách sinh viên trong lớp]"
    },
    {
        "name": "Màn hình quản lý phòng học",
        "type": "danh mục",
        "member": MEMBERS[1],
        "route": "/admin/rooms",
        "view": "src/views/pages/admin/rooms.pug",
        "js": "src/public/js/admin/rooms.js",
        "function": "Quản lý phòng học theo học kỳ, loại phòng, sức chứa, trạng thái và lớp đang dùng.",
        "layout": "Màn hình có tìm kiếm phòng, lọc học kỳ, loại phòng, tình trạng dùng, trạng thái; bảng phòng; modal phòng; modal lớp đang dùng.",
        "objects": rows(
            obj("Bộ lọc phòng học", "select group", "Tìm theo mã, tên, tòa nhà, loại phòng hoặc học kỳ", "Lọc phòng học."),
            obj("Bảng phòng học", "table", "Cột mã, tên, tòa nhà, sức chứa, loại, học kỳ, trạng thái", "Hiển thị phòng học theo học kỳ."),
            obj("Form phòng học", "modal form", "Mã phòng, tên phòng, học kỳ bắt buộc, sức chứa lớn hơn 0", "Thêm hoặc sửa phòng học."),
            obj("Modal lớp đang dùng", "modal table", "Lọc theo học kỳ", "Xem các lớp đang sử dụng phòng.")
        ),
        "events": rows(
            evt("Thêm hoặc sửa phòng", "Gửi POST hoặc PUT /api/rooms với thông tin phòng và học kỳ áp dụng."),
            evt("Xem lớp đang dùng", "Gọi /api/rooms/:id/classes theo học kỳ và render modal."),
            evt("Xóa phòng", "Gọi DELETE /api/rooms/:id và xử lý ràng buộc lịch học."),
            evt("Lọc phòng trống hoặc đang dùng", "Cập nhật query lọc usedStatus để hiển thị phòng phù hợp.")
        ),
        "install": "Cài rooms.pug, rooms.js, roomRoutes và roomController cho phòng học theo học kỳ.",
        "tests": "Kiểm thử thêm, sửa, xóa phòng, lọc học kỳ, lọc loại phòng, xem lớp đang dùng và lỗi phòng đang có lịch.",
        "completion": "95%",
        "proof": "[Chụp thêm phòng học thành công], [Chụp lớp đang dùng phòng], [Chụp lỗi xóa phòng đang dùng]"
    },
    {
        "name": "Màn hình quản lý giảng viên",
        "type": "danh mục",
        "member": MEMBERS[1],
        "route": "/admin/lecturers",
        "view": "src/views/pages/admin/lecturers.pug",
        "js": "src/public/js/admin/lecturers.js",
        "function": "Quản lý giảng viên theo học kỳ, học hàm, học vị, khoa, email và lớp giảng dạy.",
        "layout": "Màn hình có tìm kiếm, lọc học kỳ, khoa, trạng thái, bảng giảng viên, modal giảng viên và modal lớp giảng dạy.",
        "objects": rows(
            obj("Bộ lọc giảng viên", "select + input", "Tìm theo mã GV, họ tên, khoa, email hoặc học kỳ", "Lọc danh sách giảng viên."),
            obj("Bảng giảng viên", "table", "Cột mã GV, học hàm, học vị, họ tên, khoa, liên hệ, học kỳ, trạng thái", "Hiển thị giảng viên."),
            obj("Form giảng viên", "modal form", "Mã, họ tên, email và học kỳ áp dụng bắt buộc", "Thêm hoặc sửa giảng viên."),
            obj("Modal lớp giảng dạy", "modal table", "Lọc theo học kỳ", "Xem các lớp giảng viên phụ trách.")
        ),
        "events": rows(
            evt("Thêm hoặc sửa giảng viên", "Gửi POST hoặc PUT /api/lecturers, kiểm tra email hợp lệ."),
            evt("Xem lớp giảng dạy", "Gọi /api/lecturers/:id/classes theo học kỳ."),
            evt("Xóa giảng viên", "Gọi DELETE /api/lecturers/:id và xử lý ràng buộc lớp học."),
            evt("Lọc theo học kỳ", "Chỉ hiển thị giảng viên áp dụng trong học kỳ được chọn.")
        ),
        "install": "Cài lecturers.pug, lecturers.js, lecturerRoutes và lecturerController.",
        "tests": "Kiểm thử thêm, sửa, xóa giảng viên, lọc học kỳ, xem lớp giảng dạy, lỗi email và lỗi giảng viên đang có lớp.",
        "completion": "95%",
        "proof": "[Chụp thêm giảng viên thành công], [Chụp lớp giảng dạy], [Chụp lỗi email giảng viên]"
    },
    {
        "name": "Màn hình quản lý ràng buộc môn học",
        "type": "nghiệp vụ",
        "member": MEMBERS[1],
        "route": "/admin/prerequisites",
        "view": "src/views/pages/admin/prerequisites.pug",
        "js": "src/public/js/admin/prerequisites.js",
        "function": "Quản lý môn tiên quyết và môn học trước bằng popup chọn môn.",
        "layout": "Màn hình có tìm kiếm theo mã, tên môn hoặc tên môn điều kiện, lọc loại điều kiện, bảng ràng buộc, modal thêm sửa và popup chọn môn.",
        "objects": rows(
            obj("Bộ tìm kiếm ràng buộc", "select + input", "Tìm theo môn học hoặc môn điều kiện", "Lọc danh sách ràng buộc."),
            obj("Bảng ràng buộc", "table", "Cột môn học, môn điều kiện, loại, mô tả, trạng thái, sửa bởi, sửa lúc", "Hiển thị điều kiện học phần."),
            obj("Popup chọn môn", "modal picker", "Tìm theo mã hoặc tên môn", "Chọn môn học và môn điều kiện."),
            obj("Form ràng buộc", "modal form", "Môn học, môn điều kiện và loại điều kiện bắt buộc, không cho tự phụ thuộc", "Thêm hoặc sửa ràng buộc.")
        ),
        "events": rows(
            evt("Chọn môn điều kiện", "Gọi /api/courses và đưa mã môn vào form."),
            evt("Thêm ràng buộc", "Gửi POST /api/prerequisites, kiểm tra trùng cặp môn và loại điều kiện."),
            evt("Sửa ràng buộc", "Gửi PUT /api/prerequisites/:id."),
            evt("Xóa ràng buộc", "Gọi DELETE /api/prerequisites/:id và cập nhật bảng.")
        ),
        "install": "Cài prerequisites.pug, prerequisites.js, prerequisiteRoutes và prerequisiteController.",
        "tests": "Kiểm thử thêm, sửa, xóa ràng buộc, popup chọn môn, tìm kiếm, lỗi trùng ràng buộc và lỗi tự phụ thuộc.",
        "completion": "95%",
        "proof": "[Chụp thêm ràng buộc thành công], [Chụp popup chọn môn điều kiện], [Chụp lỗi trùng ràng buộc]"
    },
    {
        "name": "Màn hình quản lý chương trình học",
        "type": "nghiệp vụ",
        "member": MEMBERS[1],
        "route": "/admin/curriculum-programs",
        "view": "src/views/pages/admin/curriculum-programs.pug",
        "js": "src/public/js/admin/curriculum-programs.js",
        "function": "Quản lý môn trong chương trình đào tạo theo ngành và học kỳ dự kiến.",
        "layout": "Màn hình nhóm chương trình theo ngành, có bộ lọc ngành, trạng thái, tìm kiếm mã ngành, tên ngành, mã môn, tên môn; modal thêm môn và popup chọn môn.",
        "objects": rows(
            obj("Bộ lọc CTĐT", "select + input", "Tìm theo ngành hoặc môn học", "Lọc chương trình học."),
            obj("Nhóm chương trình theo ngành", "grouped list", "Mỗi ngành hiển thị các học kỳ dự kiến", "Quan sát cấu trúc chương trình."),
            obj("Form môn trong CTĐT", "modal form", "Ngành, môn học và học kỳ dự kiến bắt buộc", "Thêm hoặc sửa môn trong chương trình."),
            obj("Popup chọn môn", "modal picker", "Tìm mã hoặc tên môn học", "Chọn môn chính xác khi danh mục môn lớn.")
        ),
        "events": rows(
            evt("Thêm môn vào chương trình", "Gửi POST /api/majors/curriculum/items với ngành, môn, học kỳ dự kiến."),
            evt("Sửa môn trong chương trình", "Gửi PUT /api/majors/curriculum/items/:itemId."),
            evt("Xóa môn khỏi chương trình", "Gọi DELETE /api/majors/curriculum/items/:itemId."),
            evt("Lọc theo ngành hoặc trạng thái", "Cập nhật query và render lại các nhóm chương trình.")
        ),
        "install": "Cài curriculum-programs.pug, curriculum-programs.js và API /api/majors/curriculum/items.",
        "tests": "Kiểm thử thêm, sửa, xóa môn CTĐT, lọc theo ngành, popup chọn môn, lỗi trùng môn trong ngành và kiểm tra học kỳ dự kiến.",
        "completion": "95%",
        "proof": "[Chụp thêm môn vào CTĐT], [Chụp popup chọn môn CTĐT], [Chụp lỗi trùng môn CTĐT]"
    },
    {
        "name": "Màn hình quản lý môn đã học",
        "type": "nghiệp vụ",
        "member": MEMBERS[1],
        "route": "/admin/completed-courses",
        "view": "src/views/pages/admin/completed-courses.pug",
        "js": "src/public/js/admin/completed-courses.js",
        "function": "Quản lý kết quả môn đã học của sinh viên, import Excel và xem chi tiết theo từng sinh viên.",
        "layout": "Màn hình có tìm kiếm theo MSSV, họ tên, học kỳ, lọc học kỳ và kết quả, bảng tổng hợp sinh viên, modal chi tiết, modal thêm sửa và modal nhập Excel.",
        "objects": rows(
            obj("Tìm kiếm môn đã học", "select + input", "Tìm theo MSSV, họ tên hoặc học kỳ", "Lọc sinh viên có môn đã học."),
            obj("Bảng tổng hợp sinh viên", "table", "Cột MSSV, họ tên, số lượt học, qua môn, rớt, tín chỉ tích lũy", "Tổng hợp kết quả học tập."),
            obj("Form môn đã học", "modal form", "MSSV, môn học, học kỳ, lần học và kết quả bắt buộc", "Thêm hoặc sửa kết quả môn đã học."),
            obj("Modal nhập Excel", "file + table", "File .xlsx đúng mẫu gồm MSSV, MaMonHoc, HocKy, KetQua", "Import kết quả và xem dòng lỗi.")
        ),
        "events": rows(
            evt("Xem chi tiết môn đã học", "Gọi /api/completed-courses?MaSv=... và mở modal chi tiết."),
            evt("Thêm hoặc sửa kết quả", "Gọi POST hoặc PUT /api/completed-courses."),
            evt("Nhập Excel", "Gửi file đến /api/completed-courses/import và hiển thị kết quả import."),
            evt("Xóa kết quả", "Gọi DELETE /api/completed-courses/:id.")
        ),
        "install": "Cài completed-courses.pug, completed-courses.js, completedCourseRoutes và completedCourseController.",
        "tests": "Kiểm thử thêm, sửa, xóa môn đã học, xem chi tiết sinh viên, import Excel thành công, import lỗi và lỗi trùng lần học.",
        "completion": "95%",
        "proof": "[Chụp thêm môn đã học], [Chụp import môn đã học], [Chụp lỗi trùng môn đã học]"
    },
    {
        "name": "Màn hình quản lý đăng ký môn học",
        "type": "nghiệp vụ",
        "member": MEMBERS[1],
        "route": "/admin/registrations",
        "view": "src/views/pages/admin/registrations.pug",
        "js": "src/public/js/admin/registrations.js",
        "function": "Theo dõi phiếu đăng ký học phần theo sinh viên, học kỳ, trạng thái và hỗ trợ chốt đăng ký.",
        "layout": "Màn hình có vùng thời gian đăng ký, tìm kiếm theo MSSV, họ tên hoặc số phiếu, lọc học kỳ, lọc trạng thái, bảng phiếu đăng ký và modal chi tiết.",
        "objects": rows(
            obj("Bộ lọc đăng ký", "select + input", "Tìm theo tiêu chí và lọc học kỳ, trạng thái", "Tra cứu phiếu đăng ký."),
            obj("Thời gian đăng ký", "status panel", "Dữ liệu từ học kỳ đang chọn", "Cho biết khả năng chốt đăng ký."),
            obj("Bảng phiếu đăng ký", "table", "Cột MSSV, họ tên, ngành khoa, số phiếu, học kỳ, môn đã đăng ký, tín chỉ, trạng thái", "Theo dõi đăng ký toàn hệ thống."),
            obj("Modal chi tiết", "modal table", "Chỉ hiển thị môn trong học kỳ của phiếu được chọn", "Xem chi tiết học phần của sinh viên.")
        ),
        "events": rows(
            evt("Tìm kiếm phiếu đăng ký", "Cập nhật query và render danh sách đăng ký."),
            evt("Xem chi tiết", "Gọi /api/registrations/student/:studentId theo học kỳ và hiển thị modal."),
            evt("Chốt đăng ký", "Gọi /api/semesters/:id/finalize-registration khi đủ điều kiện."),
            evt("Hủy chốt đăng ký", "Gọi /api/semesters/:id/cancel-finalize-registration khi cần mở lại.")
        ),
        "install": "Cài registrations.pug, registrations.js, registrationRoutes và semester workflow endpoints.",
        "tests": "Kiểm thử tìm kiếm, lọc học kỳ, xem chi tiết đúng học kỳ, chốt đăng ký, hủy chốt và lỗi còn đơn cứu xét chờ duyệt.",
        "completion": "90%",
        "proof": "[Chụp chi tiết phiếu đăng ký], [Chụp chốt đăng ký thành công], [Chụp lỗi chốt khi còn đơn chờ]"
    },
    {
        "name": "Màn hình quản lý đơn cứu xét",
        "type": "nghiệp vụ",
        "member": MEMBERS[1],
        "route": "/admin/appeals",
        "view": "src/views/pages/admin/appeals.pug",
        "js": "src/public/js/admin/appeals.js",
        "function": "Tra cứu, duyệt hoặc từ chối đơn cứu xét đăng ký học phần của sinh viên.",
        "layout": "Màn hình có bộ lọc tìm kiếm, học kỳ, loại đơn, trạng thái, vùng thời gian cứu xét và bảng danh sách đơn.",
        "objects": rows(
            obj("Bộ lọc đơn", "select + input", "Tìm theo MSSV, họ tên, mã đơn hoặc nội dung", "Tra cứu đơn cứu xét."),
            obj("Thời gian cứu xét", "status panel", "Dựa trên học kỳ được chọn", "Cho biết thời gian xử lý đơn."),
            obj("Bảng đơn cứu xét", "table", "Cột mã đơn, sinh viên, học kỳ, loại đơn, nội dung, lý do, trạng thái", "Hiển thị đơn cần xử lý."),
            obj("Nút duyệt và từ chối", "button group", "Chỉ hiện khi đơn còn chờ duyệt", "Xử lý kết quả cứu xét.")
        ),
        "events": rows(
            evt("Tải danh sách đơn", "Gọi /api/appeals với bộ lọc hiện tại."),
            evt("Duyệt đơn", "Gọi PUT /api/appeals/:id/approve và hiển thị lỗi chi tiết nếu nghiệp vụ không hợp lệ."),
            evt("Từ chối đơn", "Nhập lý do từ chối và gọi PUT /api/appeals/:id/reject."),
            evt("Reset bộ lọc", "Xóa bộ lọc và tải lại danh sách đơn.")
        ),
        "install": "Cài appeals.pug, appeals.js, appealRoutes và appealController.",
        "tests": "Kiểm thử tra cứu đơn, lọc loại đơn, lọc trạng thái, duyệt đơn, từ chối đơn và lỗi duyệt do trùng lịch hoặc hết chỗ.",
        "completion": "90%",
        "proof": "[Chụp duyệt đơn thành công], [Chụp từ chối đơn], [Chụp lỗi duyệt đơn chi tiết]"
    },
    {
        "name": "Màn hình bảng điều khiển sinh viên",
        "type": "sinh viên",
        "member": MEMBERS[2],
        "route": "/student/dashboard",
        "view": "src/views/pages/student/dashboard.pug",
        "js": "src/public/js/student/dashboard.js",
        "function": "Tổng quan học tập, công nợ, lịch học hôm nay, thông báo mới và truy cập nhanh của sinh viên.",
        "layout": "Màn hình student có lời chào, thẻ tín chỉ hoàn thành, công nợ, lịch hôm nay, thông báo mới, danh sách thông báo và lối tắt chức năng.",
        "objects": rows(
            obj("Thẻ tín chỉ hoàn thành", "stat card", "Tính từ /api/completed-courses/me", "Hiển thị tiến độ học tập."),
            obj("Thẻ công nợ", "stat card", "Tính từ /api/tuition/student/:id", "Hiển thị học phí còn lại."),
            obj("Thẻ lịch hôm nay", "stat card", "Tính từ các đăng ký đang hiệu lực", "Nhắc lịch học trong ngày."),
            obj("Thông báo mới", "list", "Lấy từ /api/notifications", "Hiển thị các thông báo gần nhất.")
        ),
        "events": rows(
            evt("Tải dashboard sinh viên", "Gọi /api/auth/me, /api/registrations/student/:id, /api/tuition/student/:id, /api/completed-courses/me và /api/notifications."),
            evt("Có thông báo chưa đọc", "Gọi /api/notifications/unread-count và cập nhật thẻ số liệu."),
            evt("Không có dữ liệu học phí", "Hiển thị công nợ bằng 0 và giữ giao diện ổn định."),
            evt("Chọn lối tắt", "Điều hướng sang đăng ký học phần, học phí, phiếu đăng ký hoặc thông báo.")
        ),
        "install": "Cài student dashboard.pug, dashboard.js và các API đọc dữ liệu cá nhân.",
        "tests": "Kiểm thử tải dashboard, thẻ công nợ, lịch hôm nay, thông báo mới, trạng thái không dữ liệu và các lối tắt.",
        "completion": "95%",
        "proof": "[Chụp dashboard sinh viên], [Chụp thông báo mới], [Chụp thẻ công nợ]"
    },
    {
        "name": "Màn hình đăng ký học phần",
        "type": "sinh viên",
        "member": MEMBERS[2],
        "route": "/student/course-registration",
        "view": "src/views/pages/student/course-registration.pug",
        "js": "src/public/js/student/course-registration.js",
        "function": "Sinh viên tra cứu học phần mở, lọc lịch học, đăng ký học phần hoặc gửi đơn cứu xét thêm học phần.",
        "layout": "Màn hình gồm học kỳ đăng ký, bộ lọc môn, loại học phần, thứ, tiết bắt đầu, tiết kết thúc, vùng thời gian đăng ký và danh sách học phần có thể đăng ký.",
        "objects": rows(
            obj("Học kỳ đăng ký", "hidden + label", "Lấy từ /api/semesters/registration-options", "Xác định kỳ đăng ký hiện hành."),
            obj("Bộ lọc học phần", "select + input", "Tìm theo mã, tên, lớp, giảng viên và lọc lịch", "Tra cứu học phần mở."),
            obj("Danh sách học phần", "dynamic table", "Chỉ hiển thị lớp còn điều kiện đăng ký hoặc cứu xét", "Cho sinh viên chọn học phần."),
            obj("Modal đơn cứu xét thêm", "modal form", "Lý do tối đa 500 ký tự", "Gửi đơn thêm học phần ngoài cửa sổ đăng ký.")
        ),
        "events": rows(
            evt("Tải học phần có thể đăng ký", "Gọi /api/registrations/available với học kỳ và bộ lọc."),
            evt("Nhấn Đăng ký", "Gửi POST /api/registrations, kiểm tra học kỳ, trùng môn, trùng lịch, tiên quyết và sĩ số."),
            evt("Gửi đơn thêm", "Gửi POST /api/appeals với loại đơn thêm và lý do cứu xét."),
            evt("Reset bộ lọc", "Xóa các điều kiện tìm kiếm và tải lại danh sách môn.")
        ),
        "install": "Cài course-registration.pug, course-registration.js, registrationRoutes, appealRoutes và periodRoutes.",
        "tests": "Kiểm thử đăng ký hợp lệ, chặn trùng môn, chặn trùng lịch, hết hạn đăng ký, gửi đơn cứu xét và lọc theo tiết.",
        "completion": "90%",
        "proof": "[Chụp đăng ký học phần thành công], [Chụp lỗi trùng lịch], [Chụp gửi đơn cứu xét]"
    },
    {
        "name": "Màn hình phiếu đăng ký học phần",
        "type": "sinh viên",
        "member": MEMBERS[2],
        "route": "/student/my-courses",
        "view": "src/views/pages/student/my-courses.pug",
        "js": "src/public/js/student/my-courses.js",
        "function": "Sinh viên xem phiếu đăng ký, hủy hoặc gửi đơn cứu xét đổi lớp, hủy lớp và theo dõi đơn của mình.",
        "layout": "Màn hình có bộ lọc năm học, học kỳ, bảng chi tiết phiếu đăng ký, bảng đơn cứu xét, modal gửi đơn và modal xác nhận.",
        "objects": rows(
            obj("Bộ lọc năm học học kỳ", "select group", "Chọn năm học và học kỳ", "Lọc phiếu đăng ký của sinh viên."),
            obj("Bảng học phần đã đăng ký", "table", "Cột phiếu ĐK, mã lớp, học phần, tín chỉ, loại đăng ký, giảng viên, lịch, phòng, trạng thái", "Hiển thị chi tiết phiếu."),
            obj("Bảng đơn cứu xét của tôi", "table", "Cột mã đơn, học kỳ, loại đơn, nội dung, trạng thái, lý do", "Theo dõi đơn đã gửi."),
            obj("Modal xác nhận", "modal dialog", "Yêu cầu xác nhận trước khi hủy", "Ngăn thao tác hủy nhầm.")
        ),
        "events": rows(
            evt("Tải phiếu đăng ký", "Gọi /api/registrations/student/:id theo học kỳ."),
            evt("Hủy học phần", "Gọi PUT /api/registrations/:id/cancel sau khi xác nhận."),
            evt("Gửi đơn cứu xét đổi hoặc hủy", "Gửi POST /api/appeals với lớp hiện tại, lớp muốn đổi và lý do."),
            evt("Hủy đơn cứu xét", "Gọi PUT /api/appeals/:id/cancel nếu đơn còn chờ duyệt.")
        ),
        "install": "Cài my-courses.pug, my-courses.js, registrationRoutes và appealRoutes cho phiếu đăng ký sinh viên.",
        "tests": "Kiểm thử lọc theo học kỳ, hủy học phần, gửi đơn đổi lớp, gửi đơn hủy lớp, hủy đơn và cập nhật bảng sau thao tác.",
        "completion": "90%",
        "proof": "[Chụp phiếu đăng ký], [Chụp hủy học phần thành công], [Chụp đơn cứu xét của tôi]"
    },
    {
        "name": "Màn hình môn đã học của sinh viên",
        "type": "sinh viên",
        "member": MEMBERS[2],
        "route": "/student/completed-courses",
        "view": "src/views/pages/student/completed-courses.pug",
        "js": "src/public/js/student/completed-courses.js",
        "function": "Sinh viên tra cứu lịch sử môn đã học, lọc theo loại, khoa, học kỳ và kết quả.",
        "layout": "Màn hình có ô tìm kiếm môn, bộ lọc loại môn, khoa, học kỳ, kết quả và bảng lịch sử môn đã học.",
        "objects": rows(
            obj("Tìm kiếm môn đã học", "input text", "Tìm mã môn, tên môn, khoa, loại, lớp, giảng viên", "Lọc lịch sử học tập."),
            obj("Bộ lọc loại, khoa, học kỳ, kết quả", "select group", "Các select lấy dữ liệu từ API", "Thu hẹp danh sách môn đã học."),
            obj("Bảng lịch sử", "table", "Cột mã môn, tên môn, khoa, loại, tín chỉ, học kỳ, lớp, lần học, kết quả", "Hiển thị kết quả học tập."),
            obj("Tổng kết trạng thái", "summary area", "Tính từ danh sách đã lọc", "Cho biết số môn qua hoặc rớt.")
        ),
        "events": rows(
            evt("Tải môn đã học", "Gọi /api/completed-courses/me với bộ lọc hiện tại."),
            evt("Lọc theo khoa", "Gọi /api/faculties để đổ dữ liệu select khoa và lọc bảng."),
            evt("Đổi kết quả lọc", "Tải lại trang dữ liệu từ trang 1."),
            evt("Không có dữ liệu", "Ẩn bảng và hiển thị thông báo không tìm thấy môn.")
        ),
        "install": "Cài completed-courses.pug và student/completed-courses.js sử dụng /api/completed-courses/me.",
        "tests": "Kiểm thử tìm kiếm môn, lọc học kỳ, lọc kết quả qua rớt, lọc khoa và trạng thái không có dữ liệu.",
        "completion": "95%",
        "proof": "[Chụp lịch sử môn đã học], [Chụp lọc kết quả rớt], [Chụp không có dữ liệu]"
    },
    {
        "name": "Màn hình chương trình đào tạo sinh viên",
        "type": "sinh viên",
        "member": MEMBERS[2],
        "route": "/student/curriculum",
        "view": "src/views/pages/student/curriculum.pug",
        "js": "src/public/js/student/curriculum.js",
        "function": "Sinh viên xem chương trình đào tạo của ngành và tiến độ tích lũy.",
        "layout": "Màn hình có tiêu đề chương trình đào tạo, thẻ tiến độ tích lũy, bộ lọc trạng thái và danh sách môn trong chương trình.",
        "objects": rows(
            obj("Tiến độ tích lũy", "progress summary", "Tính theo môn đã hoàn thành", "Hiển thị mức độ hoàn thành CTĐT."),
            obj("Bộ lọc trạng thái", "select", "Lọc tất cả, đã học, chưa học hoặc đang học", "Theo dõi môn trong chương trình."),
            obj("Danh sách môn CTĐT", "dynamic list", "Dữ liệu từ /api/courses/curriculum/me", "Hiển thị môn theo học kỳ dự kiến."),
            obj("Trạng thái môn", "badge", "Dựa trên kết quả học tập", "Phân biệt môn đã hoàn thành và chưa hoàn thành.")
        ),
        "events": rows(
            evt("Tải chương trình đào tạo", "Gọi /api/courses/curriculum/me và lưu danh sách môn."),
            evt("Đổi trạng thái lọc", "Render lại danh sách theo status-filter."),
            evt("Tính tiến độ", "Cập nhật tổng số tín chỉ và số tín chỉ đã tích lũy."),
            evt("API lỗi", "Hiển thị trạng thái lỗi tải chương trình đào tạo.")
        ),
        "install": "Cài curriculum.pug và student/curriculum.js dùng API /api/courses/curriculum/me.",
        "tests": "Kiểm thử tải CTĐT, lọc môn đã học, lọc môn chưa học, tính tiến độ tín chỉ và xử lý khi chưa có dữ liệu.",
        "completion": "95%",
        "proof": "[Chụp CTĐT sinh viên], [Chụp lọc môn đã học], [Chụp tiến độ tích lũy]"
    },
    {
        "name": "Màn hình thời khóa biểu sinh viên",
        "type": "sinh viên",
        "member": MEMBERS[2],
        "route": "/student/my-schedule",
        "view": "src/views/pages/student/my-schedule.pug",
        "js": "src/public/js/student/my-schedule.js",
        "function": "Sinh viên xem lịch học trong tuần theo học kỳ từ các học phần đã đăng ký.",
        "layout": "Màn hình có chọn học kỳ, vùng lịch học trong tuần, mỗi lịch hiển thị học phần, giảng viên, tiết học và phòng.",
        "objects": rows(
            obj("Chọn học kỳ", "select", "Được đổ từ danh sách đăng ký của sinh viên", "Lọc lịch học theo học kỳ."),
            obj("Lịch học trong tuần", "schedule grid", "Nhóm theo thứ và sắp xếp theo tiết", "Hiển thị thời khóa biểu."),
            obj("Thẻ lịch học", "schedule item", "Có môn, lớp, giảng viên, phòng, tiết", "Cho sinh viên xem chi tiết buổi học."),
            obj("URL đồng bộ học kỳ", "query state", "Lưu học kỳ đang chọn trên URL", "Giữ trạng thái khi tải lại.")
        ),
        "events": rows(
            evt("Tải lịch học", "Gọi /api/auth/me và /api/registrations/student/:id để lấy lớp đã đăng ký."),
            evt("Đổi học kỳ", "Lọc lại danh sách lớp và cập nhật URL."),
            evt("Có lịch cũ dạng text", "Parse lịch học legacy để hiển thị được thứ, tiết và phòng."),
            evt("Không có lịch", "Hiển thị thông báo chưa có lịch học cho học kỳ.")
        ),
        "install": "Cài my-schedule.pug và my-schedule.js dùng dữ liệu đăng ký hiện hành.",
        "tests": "Kiểm thử tải lịch học, đổi học kỳ, hiển thị phòng học, sắp xếp theo tiết và trạng thái chưa có lịch.",
        "completion": "95%",
        "proof": "[Chụp thời khóa biểu], [Chụp đổi học kỳ lịch học], [Chụp trạng thái chưa có lịch]"
    },
    {
        "name": "Màn hình học phí của sinh viên",
        "type": "sinh viên",
        "member": MEMBERS[2],
        "route": "/student/my-tuition",
        "view": "src/views/pages/student/my-tuition.pug",
        "js": "src/public/js/student/my-tuition.js",
        "function": "Sinh viên xem học phí theo học kỳ, phiếu thu còn nợ và tạo thanh toán trực tuyến hoặc ghi nhận phương thức thanh toán.",
        "layout": "Màn hình có thẻ tổng học phí, đã đóng, còn lại, bảng học phí theo học kỳ, modal thanh toán, modal chi tiết học phí và modal phiếu thu.",
        "objects": rows(
            obj("Thẻ tổng học phí", "stat cards", "Tính từ /api/tuition/student/:id", "Tổng hợp học phí, đã đóng và còn lại."),
            obj("Bảng học phí", "table", "Cột học kỳ, tổng học phí, đã đóng, còn lại, hạn đóng, trạng thái", "Hiển thị công nợ theo học kỳ."),
            obj("Modal thanh toán", "modal form", "Số tiền lớn hơn 0, phương thức hợp lệ", "Tạo thanh toán cho phiếu thu còn nợ."),
            obj("Modal chi tiết", "modal detail", "Lấy dữ liệu theo phiếu đăng ký", "Xem học phần và tiền học phí chi tiết.")
        ),
        "events": rows(
            evt("Tải học phí", "Gọi /api/auth/me và /api/tuition/student/:id."),
            evt("Mở thanh toán", "Chọn phiếu thu còn nợ và mở modal nhập số tiền, phương thức."),
            evt("Tạo thanh toán", "Gọi POST /api/payments/:id/checkout hoặc /api/payments/checkout."),
            evt("Xem chi tiết học phí", "Gọi /api/tuition/detail/:id và hiển thị danh sách học phần.")
        ),
        "install": "Cài my-tuition.pug, my-tuition.js, tuitionRoutes, paymentRoutes và payment gateway config.",
        "tests": "Kiểm thử xem học phí, xem chi tiết, tạo thanh toán tiền mặt, tạo thanh toán online, lỗi số tiền không hợp lệ và cập nhật trạng thái.",
        "completion": "90%",
        "proof": "[Chụp học phí sinh viên], [Chụp thanh toán thành công], [Chụp lỗi số tiền thanh toán]"
    },
    {
        "name": "Màn hình phiếu thu của sinh viên",
        "type": "sinh viên",
        "member": MEMBERS[2],
        "route": "/student/my-payments",
        "view": "src/views/pages/student/my-payments.pug",
        "js": "src/public/js/student/my-payments.js",
        "function": "Sinh viên xem danh sách phiếu thu, trạng thái thanh toán và chi tiết từng phiếu.",
        "layout": "Màn hình có bộ lọc học kỳ, bảng phiếu thu, modal chi tiết phiếu thu với số tiền, phương thức, mã giao dịch, xác nhận và trạng thái.",
        "objects": rows(
            obj("Lọc học kỳ", "select", "Lấy danh sách học kỳ từ API", "Lọc phiếu thu theo học kỳ."),
            obj("Bảng phiếu thu", "table", "Cột số phiếu, học kỳ, số tiền, phương thức, ngày lập, mã giao dịch, xác nhận, trạng thái", "Hiển thị phiếu thu của sinh viên."),
            obj("Modal chi tiết phiếu thu", "modal detail", "Lấy theo số phiếu thu", "Xem đầy đủ thông tin phiếu thu."),
            obj("Trạng thái phiếu", "badge", "Chuẩn hóa trạng thái từ API", "Phân biệt chưa thanh toán, thành công, thất bại, hủy.")
        ),
        "events": rows(
            evt("Tải phiếu thu", "Gọi /api/auth/me, /api/semesters và /api/payments/student/:id."),
            evt("Đổi học kỳ", "Gọi loadMyPayments từ trang 1."),
            evt("Xem chi tiết phiếu", "Gọi /api/payments/:id và mở modal chi tiết."),
            evt("Không có phiếu", "Ẩn bảng và hiển thị thông báo chưa có phiếu thu.")
        ),
        "install": "Cài my-payments.pug và my-payments.js dùng paymentRoutes.",
        "tests": "Kiểm thử lọc học kỳ, xem chi tiết phiếu thu, trạng thái thành công, trạng thái thất bại và trường hợp chưa có phiếu.",
        "completion": "95%",
        "proof": "[Chụp danh sách phiếu thu sinh viên], [Chụp chi tiết phiếu thu], [Chụp trạng thái thất bại]"
    },
    {
        "name": "Màn hình thông báo sinh viên",
        "type": "sinh viên",
        "member": MEMBERS[2],
        "route": "/student/notifications",
        "view": "src/views/pages/student/notifications.pug",
        "js": "src/public/js/student/notifications.js",
        "function": "Sinh viên xem thông báo, lọc đã đọc chưa đọc và mở chi tiết thông báo.",
        "layout": "Màn hình có bộ lọc trạng thái đọc, danh sách tin mới từ hệ thống và modal chi tiết thông báo.",
        "objects": rows(
            obj("Lọc trạng thái đọc", "select", "Tất cả, đã đọc hoặc chưa đọc", "Lọc danh sách thông báo."),
            obj("Danh sách thông báo", "list", "Lấy từ /api/notifications", "Hiển thị tiêu đề, nội dung tóm tắt, ngày tạo."),
            obj("Modal chi tiết", "modal detail", "Mở theo mã thông báo", "Đọc đầy đủ nội dung thông báo."),
            obj("Bộ đếm chưa đọc", "counter", "Dữ liệu từ /api/notifications/unread-count", "Cập nhật số thông báo mới.")
        ),
        "events": rows(
            evt("Tải thông báo", "Gọi /api/notifications với bộ lọc trạng thái đọc."),
            evt("Mở chi tiết thông báo", "Gọi /api/notifications/:id, hiển thị modal và đánh dấu đã đọc."),
            evt("Đánh dấu đã đọc", "Gọi PUT /api/notifications/:id/read."),
            evt("Đổi bộ lọc", "Tải lại danh sách từ trang 1.")
        ),
        "install": "Cài notifications.pug và student/notifications.js sử dụng notificationRoutes.",
        "tests": "Kiểm thử xem danh sách, lọc chưa đọc, mở chi tiết, đánh dấu đã đọc và cập nhật số thông báo.",
        "completion": "95%",
        "proof": "[Chụp danh sách thông báo], [Chụp chi tiết thông báo], [Chụp lọc chưa đọc]"
    },
    {
        "name": "Màn hình hồ sơ sinh viên",
        "type": "sinh viên",
        "member": MEMBERS[2],
        "route": "/student/profile",
        "view": "src/views/pages/student/profile.pug",
        "js": "src/public/js/student/profile.js",
        "function": "Sinh viên xem hồ sơ học vụ, cập nhật số điện thoại, địa chỉ, giới tính và ảnh đại diện.",
        "layout": "Màn hình có thông tin học vụ, ảnh đại diện, các trường khóa chỉ xem và các trường cho phép chỉnh sửa.",
        "objects": rows(
            obj("Ảnh đại diện", "file upload", "Chấp nhận jpeg, png, webp, gif", "Cập nhật avatar sinh viên."),
            obj("Trường khóa", "readonly inputs", "MSSV, họ tên, email, ngày sinh, CCCD, dân tộc, ngành, khoa, trạng thái, khóa không sửa", "Hiển thị dữ liệu học vụ."),
            obj("Trường chỉnh sửa", "input/select", "Số điện thoại, địa chỉ và giới tính được chỉnh sửa", "Cập nhật thông tin liên hệ."),
            obj("Nút Lưu thay đổi", "button", "Kiểm tra dữ liệu trước khi gửi", "Gửi cập nhật hồ sơ.")
        ),
        "events": rows(
            evt("Tải hồ sơ", "Gọi /api/auth/me và đổ dữ liệu vào form."),
            evt("Lưu hồ sơ", "Gửi PUT /api/auth/profile với các trường được phép sửa."),
            evt("Cập nhật ảnh", "Gửi POST /api/auth/avatar bằng multipart form data."),
            evt("Nhấn trường khóa", "Hiển thị thông báo trường không được sửa.")
        ),
        "install": "Cài profile.pug và student/profile.js dùng API /api/auth/me, /api/auth/profile, /api/auth/avatar.",
        "tests": "Kiểm thử tải hồ sơ, lưu số điện thoại, lưu địa chỉ, upload ảnh, lỗi file ảnh quá lớn và cảnh báo trường khóa.",
        "completion": "95%",
        "proof": "[Chụp hồ sơ sinh viên], [Chụp cập nhật hồ sơ thành công], [Chụp cảnh báo trường khóa]"
    },
    {
        "name": "Màn hình hồ sơ quản trị viên",
        "type": "hệ thống",
        "member": MEMBERS[2],
        "route": "/admin/profile",
        "view": "src/views/pages/admin/profile.pug",
        "js": "src/public/js/admin/profile.js",
        "function": "Quản trị viên xem và cập nhật hồ sơ cá nhân, ảnh đại diện, họ tên, email, số điện thoại, chức vụ, phòng ban.",
        "layout": "Màn hình admin có avatar, thông tin tài khoản, các trường hồ sơ và nút lưu hồ sơ.",
        "objects": rows(
            obj("Ảnh quản trị viên", "file upload", "File ảnh đúng định dạng", "Cập nhật ảnh đại diện."),
            obj("Tên đăng nhập", "readonly input", "Không được chỉnh sửa", "Hiển thị tài khoản đang đăng nhập."),
            obj("Thông tin hồ sơ", "input fields", "Họ tên, email, số điện thoại, chức vụ, phòng ban", "Cập nhật thông tin quản trị viên."),
            obj("Nút Lưu hồ sơ", "button", "Gửi khi form hợp lệ", "Lưu thay đổi hồ sơ.")
        ),
        "events": rows(
            evt("Tải hồ sơ admin", "Gọi /api/auth/me và đổ dữ liệu vào form."),
            evt("Lưu hồ sơ", "Gửi PUT /api/auth/profile."),
            evt("Upload avatar", "Gửi POST /api/auth/avatar và cập nhật ảnh trên header, sidebar."),
            evt("Nhập email sai định dạng", "Trình duyệt hoặc API trả lỗi và không lưu hồ sơ.")
        ),
        "install": "Cài admin profile.pug và admin/profile.js dùng authRoutes.",
        "tests": "Kiểm thử tải hồ sơ, sửa họ tên, sửa email, upload ảnh, lỗi email và trường tên đăng nhập chỉ xem.",
        "completion": "95%",
        "proof": "[Chụp hồ sơ admin], [Chụp lưu hồ sơ thành công], [Chụp upload avatar admin]"
    },
    {
        "name": "Màn hình công nợ học phí",
        "type": "nghiệp vụ",
        "member": MEMBERS[3],
        "route": "/admin/tuition",
        "view": "src/views/pages/admin/tuition.pug",
        "js": "src/public/js/admin/tuition.js",
        "function": "Quản trị viên theo dõi công nợ học phí theo sinh viên, học kỳ và trạng thái thanh toán.",
        "layout": "Màn hình có tìm kiếm MSSV hoặc họ tên, lọc học kỳ, lọc trạng thái, bảng công nợ và modal chi tiết học phí.",
        "objects": rows(
            obj("Tìm kiếm công nợ", "select + input", "Tìm theo MSSV hoặc họ tên", "Tra cứu công nợ sinh viên."),
            obj("Bộ lọc học kỳ, trạng thái", "select group", "Trạng thái chưa đóng, đóng một phần, quá hạn, đã đóng đủ, chưa phát sinh", "Lọc dữ liệu công nợ."),
            obj("Bảng công nợ", "table", "Cột MSSV, họ tên, học kỳ, số môn, tổng phải đóng, đã đóng, còn nợ, hạn đóng, trạng thái", "Theo dõi nợ học phí."),
            obj("Modal chi tiết học phí", "modal detail", "Lấy theo số phiếu", "Xem chi tiết học phần và số tiền.")
        ),
        "events": rows(
            evt("Tìm kiếm hoặc lọc công nợ", "Cập nhật query và render lại bảng công nợ."),
            evt("Xem chi tiết học phí", "Gọi /api/tuition/detail/:id và mở modal."),
            evt("Trạng thái quá hạn", "So sánh hạn đóng và số tiền đã đóng để hiển thị badge quá hạn."),
            evt("API lỗi", "Hiển thị thông báo lỗi và không làm mất bộ lọc hiện tại.")
        ),
        "install": "Cài tuition.pug, tuition.js, tuitionRoutes và tuitionController.",
        "tests": "Kiểm thử tìm theo MSSV, lọc học kỳ, lọc quá hạn, xem chi tiết học phí và dữ liệu sinh viên đã đóng đủ.",
        "completion": "95%",
        "proof": "[Chụp công nợ học phí], [Chụp chi tiết học phí], [Chụp lọc quá hạn]"
    },
    {
        "name": "Màn hình phiếu thu học phí",
        "type": "nghiệp vụ",
        "member": MEMBERS[3],
        "route": "/admin/payments",
        "view": "src/views/pages/admin/payments.pug",
        "js": "src/public/js/admin/payments.js",
        "function": "Quản lý phiếu thu học phí, tạo phiếu đơn lẻ, tạo hàng loạt, xác nhận, đánh dấu thất bại, hủy và xuất Excel.",
        "layout": "Màn hình có bộ lọc phiếu thu, bảng phiếu thu, modal lập phiếu, modal tạo hàng loạt, modal chi tiết và khối thời gian thu học phí.",
        "objects": rows(
            obj("Bộ lọc phiếu thu", "select + input", "Tìm phiếu thu, MSSV, họ tên, lọc học kỳ, phương thức, trạng thái", "Tra cứu phiếu thu."),
            obj("Bảng phiếu thu", "table", "Cột số phiếu, MSSV, họ tên, học kỳ, số tiền, phương thức, mã giao dịch, xác nhận, trạng thái", "Theo dõi thanh toán."),
            obj("Form lập phiếu thu", "modal form", "MSSV, học kỳ, số tiền và phương thức bắt buộc", "Tạo phiếu thu thủ công."),
            obj("Tạo phiếu hàng loạt", "modal action", "Chọn học kỳ", "Tạo phiếu cho các đăng ký còn nợ.")
        ),
        "events": rows(
            evt("Lập phiếu thu", "Gửi POST /api/payments."),
            evt("Tạo hàng loạt", "Gửi POST /api/payments/bulk theo học kỳ."),
            evt("Xác nhận, thất bại hoặc hủy phiếu", "Gọi PUT /api/payments/:id/confirm, /fail hoặc /cancel."),
            evt("Xuất Excel phiếu thu", "Gọi /api/payments/export với bộ lọc hiện tại.")
        ),
        "install": "Cài payments.pug, payments.js, paymentRoutes và paymentController.",
        "tests": "Kiểm thử tạo phiếu, tạo hàng loạt, xác nhận thanh toán, đánh dấu thất bại, hủy phiếu, xuất Excel và cập nhật công nợ.",
        "completion": "90%",
        "proof": "[Chụp lập phiếu thu thành công], [Chụp tạo phiếu hàng loạt], [Chụp thanh toán thành công thất bại]"
    },
    {
        "name": "Màn hình đơn giá tín chỉ",
        "type": "danh mục",
        "member": MEMBERS[3],
        "route": "/admin/pricing",
        "view": "src/views/pages/admin/pricing.pug",
        "js": "src/public/js/admin/pricing.js",
        "function": "Quản lý đơn giá tín chỉ theo loại môn, loại học, học kỳ áp dụng và trạng thái.",
        "layout": "Màn hình có tìm kiếm, lọc loại môn, loại học, học kỳ, trạng thái, bảng đơn giá và modal thêm sửa.",
        "objects": rows(
            obj("Tìm kiếm đơn giá", "select + input", "Tìm theo loại môn, loại học, đơn giá hoặc học kỳ", "Tra cứu bảng giá."),
            obj("Bảng đơn giá", "table", "Cột loại môn, loại học, đơn giá, học kỳ, trạng thái, sửa bởi, sửa lúc", "Hiển thị giá tín chỉ."),
            obj("Form đơn giá", "modal form", "Loại môn, loại học, đơn giá, học kỳ bắt buộc, đơn giá lớn hơn 0", "Thêm hoặc sửa đơn giá."),
            obj("Trạng thái áp dụng", "select", "Đang áp dụng hoặc tạm ngưng", "Điều khiển giá được dùng khi tính học phí.")
        ),
        "events": rows(
            evt("Thêm đơn giá", "Gửi POST /api/pricing."),
            evt("Sửa đơn giá", "Gửi PUT /api/pricing/:id."),
            evt("Xóa đơn giá", "Gọi DELETE /api/pricing/:id và kiểm tra phiếu đăng ký đã dùng giá."),
            evt("Click dòng", "Hiển thị chi tiết đơn giá tín chỉ.")
        ),
        "install": "Cài pricing.pug, pricing.js, pricingRoutes và pricingController.",
        "tests": "Kiểm thử thêm, sửa, xóa đơn giá, lọc học kỳ, lọc loại học, lỗi trùng loại giá trong học kỳ và lỗi đơn giá âm.",
        "completion": "95%",
        "proof": "[Chụp thêm đơn giá thành công], [Chụp lỗi trùng đơn giá], [Chụp lọc học kỳ đơn giá]"
    },
    {
        "name": "Màn hình đối tượng ưu tiên",
        "type": "danh mục",
        "member": MEMBERS[3],
        "route": "/admin/beneficiaries",
        "view": "src/views/pages/admin/beneficiaries.pug",
        "js": "src/public/js/admin/beneficiaries.js",
        "function": "Quản lý đối tượng ưu tiên, tỷ lệ giảm học phí, độ ưu tiên tự tính và danh sách sinh viên thuộc đối tượng.",
        "layout": "Màn hình có tìm kiếm theo mã hoặc tên đối tượng, bảng đối tượng, modal thêm sửa, modal sinh viên thuộc đối tượng và nhập Excel danh sách sinh viên.",
        "objects": rows(
            obj("Tìm kiếm đối tượng", "select + input", "Tìm theo mã hoặc tên đối tượng", "Lọc danh sách ưu tiên."),
            obj("Bảng đối tượng", "table", "Cột mã, tên, tỷ lệ giảm, độ ưu tiên, số SV, mô tả", "Hiển thị chính sách miễn giảm."),
            obj("Form đối tượng", "modal form", "Mã, tên và tỷ lệ giảm bắt buộc, tỷ lệ từ 0 đến 100", "Thêm hoặc sửa đối tượng."),
            obj("Modal sinh viên thuộc đối tượng", "modal table + input", "MSSV hợp lệ và không trùng", "Thêm, xóa hoặc import sinh viên thuộc đối tượng.")
        ),
        "events": rows(
            evt("Thêm hoặc sửa đối tượng", "Gọi POST hoặc PUT /api/beneficiaries, cập nhật lại độ ưu tiên theo tỷ lệ giảm."),
            evt("Thêm sinh viên vào đối tượng", "Gửi POST /api/beneficiaries/:id/students."),
            evt("Import sinh viên", "Gửi file đến /api/beneficiaries/:id/students/import."),
            evt("Xóa đối tượng hoặc sinh viên", "Gọi DELETE tương ứng và cập nhật bảng.")
        ),
        "install": "Cài beneficiaries.pug, beneficiaries.js, beneficiaryRoutes và beneficiaryController.",
        "tests": "Kiểm thử thêm, sửa, xóa đối tượng, tỷ lệ giảm, độ ưu tiên tự cập nhật, thêm SV, import SV và lỗi MSSV trùng.",
        "completion": "95%",
        "proof": "[Chụp thêm đối tượng ưu tiên], [Chụp thêm sinh viên vào đối tượng], [Chụp lỗi MSSV trùng]"
    },
    {
        "name": "Màn hình báo cáo thống kê",
        "type": "báo cáo",
        "member": MEMBERS[3],
        "route": "/admin/reports",
        "view": "src/views/pages/admin/reports.pug",
        "js": "src/public/js/admin/reports.js",
        "function": "Hiển thị thống kê sinh viên, môn học, doanh thu theo tháng, thanh toán theo phương thức và sinh viên theo trạng thái.",
        "layout": "Màn hình có chọn học kỳ, thẻ số liệu, biểu đồ doanh thu, biểu đồ thanh toán, bảng sinh viên theo trạng thái và bảng phương thức thanh toán.",
        "objects": rows(
            obj("Chọn học kỳ báo cáo", "select", "Có thể chọn toàn bộ hoặc một học kỳ", "Lọc số liệu báo cáo."),
            obj("Thẻ thống kê", "stat cards", "Lấy từ /api/students/stats và /api/courses/stats", "Hiển thị tổng quan dữ liệu."),
            obj("Biểu đồ doanh thu", "chart", "Dùng Chart.js và /api/dashboard/revenue-monthly", "Trực quan hóa doanh thu."),
            obj("Bảng thống kê", "table", "Gồm trạng thái sinh viên và phương thức thanh toán", "Đối chiếu số liệu chi tiết.")
        ),
        "events": rows(
            evt("Tải báo cáo", "Gọi loadReports để lấy thống kê sinh viên, môn học, doanh thu và thanh toán."),
            evt("Đổi học kỳ", "Thêm MaHocKy vào query thống kê và render lại biểu đồ, bảng."),
            evt("Dữ liệu rỗng", "Hiển thị số 0 và bảng rỗng có cấu trúc."),
            evt("Render biểu đồ", "Cập nhật Chart.js cho doanh thu và phương thức thanh toán.")
        ),
        "install": "Cài reports.pug, reports.js, dashboardRoutes, studentRoutes và courseRoutes cho báo cáo.",
        "tests": "Kiểm thử báo cáo toàn bộ, báo cáo theo học kỳ, biểu đồ doanh thu, thống kê phương thức thanh toán và trạng thái dữ liệu rỗng.",
        "completion": "95%",
        "proof": "[Chụp báo cáo thống kê], [Chụp biểu đồ doanh thu], [Chụp lọc theo học kỳ]"
    },
    {
        "name": "Màn hình báo cáo sinh viên chưa hoàn thành học phí",
        "type": "báo cáo",
        "member": MEMBERS[3],
        "route": "/admin/reports/incomplete-tuition",
        "view": "src/views/pages/admin/reports-incomplete-tuition.pug",
        "js": "src/public/js/admin/reports.js",
        "function": "Lập báo cáo sinh viên còn nợ hoặc quá hạn học phí, lọc theo học kỳ, khoa, ngành, trạng thái và quá hạn.",
        "layout": "Màn hình có chọn học kỳ báo cáo, nút xuất Excel, nút in, bộ lọc MSSV hoặc họ tên, khoa, ngành, trạng thái nợ và bảng báo cáo.",
        "objects": rows(
            obj("Bộ lọc báo cáo nợ", "select + input", "Học kỳ, khoa, ngành, trạng thái, quá hạn", "Lọc sinh viên chưa hoàn thành học phí."),
            obj("Bảng báo cáo nợ", "table", "Cột MSSV, họ tên, ngành khoa, học kỳ, phải đóng, đã đóng, còn nợ, hạn đóng, ngày quá hạn, trạng thái", "Hiển thị danh sách công nợ."),
            obj("Nút xuất Excel", "button", "Dùng bộ lọc hiện tại", "Tải file báo cáo."),
            obj("Nút in báo cáo", "button", "Dùng dữ liệu đang hiển thị", "In báo cáo phục vụ lưu trữ.")
        ),
        "events": rows(
            evt("Tải báo cáo nợ", "Gọi /api/dashboard/incomplete-tuition với bộ lọc."),
            evt("Lọc khoa ngành", "Khi chọn khoa, danh sách ngành được lọc tương ứng."),
            evt("Xuất Excel", "Gọi /api/dashboard/incomplete-tuition/export."),
            evt("In báo cáo", "Mở chế độ in bảng báo cáo hiện tại.")
        ),
        "install": "Cài reports-incomplete-tuition.pug và dùng reports.js với dashboardController.",
        "tests": "Kiểm thử lọc học kỳ, lọc khoa ngành, lọc quá hạn, xuất Excel, in báo cáo và kiểm tra tổng còn nợ.",
        "completion": "95%",
        "proof": "[Chụp báo cáo chưa hoàn thành học phí], [Chụp xuất Excel báo cáo nợ], [Chụp lọc quá hạn]"
    },
    {
        "name": "Màn hình quản lý người dùng",
        "type": "hệ thống",
        "member": MEMBERS[3],
        "route": "/admin/users",
        "view": "src/views/pages/admin/users.pug",
        "js": "src/public/js/admin/users.js",
        "function": "Quản lý tài khoản hệ thống, nhóm người dùng, tạo tài khoản đơn lẻ, tạo tài khoản hàng loạt cho sinh viên và xem mật khẩu tạm.",
        "layout": "Màn hình có tìm kiếm theo tên đăng nhập, họ tên, email, mã sinh viên; lọc vai trò, nhóm; bảng tài khoản; modal tạo tài khoản; modal tạo hàng loạt; modal danh sách tài khoản đã tạo.",
        "objects": rows(
            obj("Bộ lọc tài khoản", "select + input", "Tìm theo tên đăng nhập, họ tên, email, mã sinh viên", "Tra cứu người dùng."),
            obj("Bảng người dùng", "table", "Cột ID, tài khoản, liên hệ, nhóm vai trò, ngày tạo, thao tác", "Hiển thị tài khoản hệ thống."),
            obj("Form tạo tài khoản", "modal form", "Mật khẩu và xác nhận mật khẩu phải khớp", "Tạo tài khoản admin hoặc sinh viên."),
            obj("Tạo tài khoản hàng loạt", "modal form", "Chọn khoa và ngành", "Tạo tài khoản sinh viên và lưu mật khẩu tạm.")
        ),
        "events": rows(
            evt("Tạo tài khoản", "Gửi POST /api/roles/accounts sau khi kiểm tra xác nhận mật khẩu."),
            evt("Đổi nhóm tài khoản", "Gửi PUT /api/roles/accounts/:id/role."),
            evt("Tạo tài khoản sinh viên hàng loạt", "Gửi POST /api/roles/accounts/batch-create-student-accounts theo khoa ngành."),
            evt("Xem danh sách mật khẩu tạm", "Gọi /api/roles/accounts/student-credentials?limit=100 và hỗ trợ ẩn hiện mật khẩu.")
        ),
        "install": "Cài users.pug, users.js, roleRoutes và roleController.",
        "tests": "Kiểm thử tạo tài khoản, xác nhận mật khẩu, đổi nhóm, tạo hàng loạt theo khoa ngành, xem mật khẩu tạm và xóa tài khoản.",
        "completion": "90%",
        "proof": "[Chụp tạo tài khoản thành công], [Chụp lỗi xác nhận mật khẩu], [Chụp tạo tài khoản hàng loạt]"
    },
    {
        "name": "Màn hình phân quyền hệ thống",
        "type": "hệ thống",
        "member": MEMBERS[3],
        "route": "/admin/permissions",
        "view": "src/views/pages/admin/permissions.pug",
        "js": "src/public/js/admin/permissions.js",
        "function": "Quản lý nhóm người dùng, quyền truy cập và gán quyền theo nhóm.",
        "layout": "Màn hình có tab Nhóm người dùng và Quyền truy cập, bảng nhóm, bảng quyền, modal thêm sửa nhóm, modal thêm sửa quyền và modal phân quyền.",
        "objects": rows(
            obj("Tab nhóm và quyền", "segmented control", "Lưu tab đang chọn trên URL", "Chuyển giữa quản lý nhóm và quyền."),
            obj("Bảng nhóm người dùng", "table", "Cột mã nhóm, tên nhóm, loại nhóm, số tài khoản, số quyền", "Hiển thị nhóm người dùng."),
            obj("Bảng quyền truy cập", "table", "Cột mã quyền, tên quyền, loại quyền, màn hình đường dẫn", "Hiển thị chức năng được phân quyền."),
            obj("Modal phân quyền", "modal form", "Gán quyền theo mã nhóm và mã chức năng", "Cập nhật quyền truy cập màn hình.")
        ),
        "events": rows(
            evt("Thêm hoặc sửa nhóm", "Gửi POST hoặc PUT /api/permissions/groups."),
            evt("Thêm hoặc sửa quyền", "Gửi POST hoặc PUT /api/permissions/functions."),
            evt("Gán quyền", "Gửi PUT /api/permissions/groups/:id/permissions để cập nhật hàng loạt."),
            evt("Kiểm tra truy cập", "Sidebar và viewRoutes dùng quyền để ẩn hoặc chặn màn hình.")
        ),
        "install": "Cài permissions.pug, permissions.js, permissionRoutes, permissionController và permissionCatalog.",
        "tests": "Kiểm thử thêm nhóm, sửa nhóm, thêm quyền, sửa quyền, gán quyền, xóa quyền và truy cập màn hình khi không đủ quyền.",
        "completion": "90%",
        "proof": "[Chụp gán quyền thành công], [Chụp lỗi không đủ quyền], [Chụp tab quyền truy cập]"
    },
    {
        "name": "Màn hình quản lý thông báo",
        "type": "hệ thống",
        "member": MEMBERS[3],
        "route": "/admin/notifications",
        "view": "src/views/pages/admin/notifications.pug",
        "js": "src/public/js/admin/notifications.js",
        "function": "Quản trị viên tạo, sửa, xóa, lọc và xem trước thông báo gửi cho toàn hệ thống, khoa, ngành hoặc cá nhân.",
        "layout": "Màn hình có lọc nguồn và loại thông báo, bảng thông báo, modal tạo thông báo, modal xem trước, trường tiêu đề, nội dung, loại, nhóm hạn, đối tượng, khoa, ngành, ghim, ngày hết hạn, đường dẫn.",
        "objects": rows(
            obj("Bộ lọc thông báo", "select group", "Lọc theo nguồn tự động hoặc thủ công và loại thông báo", "Tra cứu thông báo."),
            obj("Bảng thông báo", "table", "Cột mã, tiêu đề, loại, nguồn, đối tượng, ghim, hết hạn, ngày tạo", "Quản lý danh sách thông báo."),
            obj("Form thông báo", "modal form", "Tiêu đề, nội dung, loại và đối tượng bắt buộc", "Tạo hoặc sửa thông báo."),
            obj("Xem trước thông báo", "modal preview", "Dùng nội dung trong form", "Kiểm tra nội dung trước khi gửi.")
        ),
        "events": rows(
            evt("Tạo thông báo", "Gửi POST /api/notifications hoặc /api/notifications/personal tùy loại đối tượng."),
            evt("Sửa thông báo", "Gửi PUT /api/notifications/:id."),
            evt("Xóa thông báo", "Gọi DELETE /api/notifications/:id."),
            evt("Lọc khoa ngành", "Khi chọn khoa, danh sách ngành nhận thông báo được lọc theo khoa.")
        ),
        "install": "Cài notifications.pug, notifications.js, notificationRoutes và notificationController.",
        "tests": "Kiểm thử tạo thông báo chung, tạo thông báo theo khoa ngành, xem trước, ghim thông báo, xóa thông báo và lọc nguồn tự động.",
        "completion": "95%",
        "proof": "[Chụp tạo thông báo], [Chụp xem trước thông báo], [Chụp lọc thông báo tự động]"
    },
    {
        "name": "Màn hình tham số hệ thống",
        "type": "hệ thống",
        "member": MEMBERS[3],
        "route": "/admin/settings",
        "view": "src/views/pages/admin/settings.pug",
        "js": "src/public/js/admin/settings.js",
        "function": "Cấu hình tham số quy chế đăng ký, ràng buộc Anh văn và điều kiện khóa luận.",
        "layout": "Màn hình chia thành các nhóm tham số: quy định tín chỉ, ràng buộc Anh văn, điều kiện khóa luận và bảng ràng buộc đang sử dụng tham số.",
        "objects": rows(
            obj("Số tín chỉ tối thiểu tối đa", "number inputs", "Giá trị nguyên không âm, tối đa phải lớn hơn tối thiểu", "Điều khiển giới hạn đăng ký."),
            obj("Danh sách môn Anh văn bắt buộc", "input text", "Danh sách mã môn phân tách bằng dấu phẩy", "Kiểm tra ràng buộc Anh văn."),
            obj("Giới hạn khóa luận", "number input", "Số tín chỉ nợ tối đa", "Kiểm tra điều kiện đăng ký khóa luận."),
            obj("Bảng tác động tham số", "table", "Liệt kê luồng bị ảnh hưởng", "Giúp admin biết thay đổi tác động đến nghiệp vụ nào.")
        ),
        "events": rows(
            evt("Cập nhật tham số", "Gửi PUT /api/settings với các giá trị đã chuẩn hóa."),
            evt("Nhập danh sách môn Anh văn", "normalizeCourseList chuẩn hóa mã môn và bỏ khoảng trắng thừa."),
            evt("Giá trị số không hợp lệ", "parseIntegerField trả lỗi và không gửi API."),
            evt("Lưu thành công", "Hiển thị thông báo và giữ dữ liệu mới trên form.")
        ),
        "install": "Cài settings.pug, settings.js, settingsRoutes và settingsController.",
        "tests": "Kiểm thử cập nhật tín chỉ, cập nhật môn Anh văn, lỗi số âm, lỗi tối đa nhỏ hơn tối thiểu và tác động đến đăng ký học phần.",
        "completion": "95%",
        "proof": "[Chụp cập nhật tham số], [Chụp lỗi tham số không hợp lệ], [Chụp bảng luồng bị ảnh hưởng]"
    },
    {
        "name": "Màn hình thùng rác",
        "type": "hệ thống",
        "member": MEMBERS[3],
        "route": "/admin/trash",
        "view": "src/views/pages/admin/trash.pug",
        "js": "src/public/js/admin/trash.js",
        "function": "Quản lý dữ liệu đã xóa mềm, khôi phục hoặc xóa vĩnh viễn theo loại dữ liệu được phân quyền.",
        "layout": "Màn hình có ô tìm kiếm, select loại thùng rác, nút làm mới, khôi phục đã chọn, xóa vĩnh viễn đã chọn và bảng bản ghi đã xóa.",
        "objects": rows(
            obj("Tìm kiếm thùng rác", "input text", "Tìm theo mã hoặc tên đã xóa", "Lọc bản ghi trong thùng rác."),
            obj("Loại thùng rác", "select", "Chỉ hiện entity người dùng có quyền", "Chọn loại dữ liệu đã xóa."),
            obj("Bảng thùng rác", "table", "Cột chọn, loại dữ liệu, khóa chính, tên hiển thị, người xóa, thời điểm xóa", "Hiển thị bản ghi đã xóa mềm."),
            obj("Nút khôi phục và xóa vĩnh viễn", "button group", "Chỉ bật khi có dòng được chọn", "Thực hiện thao tác hàng loạt hoặc từng dòng.")
        ),
        "events": rows(
            evt("Tải thùng rác", "Gọi /api/trash/:entity và render bảng."),
            evt("Khôi phục một hoặc nhiều dòng", "Gọi POST /api/trash/:entity/:id/restore hoặc /batch-restore."),
            evt("Xóa vĩnh viễn", "Gọi DELETE /api/trash/:entity/:id/purge hoặc /batch-purge sau xác nhận."),
            evt("Chọn tất cả", "Cập nhật checkbox dòng và trạng thái nút thao tác hàng loạt.")
        ),
        "install": "Cài trash.pug, trash.js, trashRoutes, trashController và trashConfig.",
        "tests": "Kiểm thử tải từng loại thùng rác, tìm kiếm, khôi phục một dòng, khôi phục hàng loạt, xóa vĩnh viễn và quyền không thấy entity bị cấm.",
        "completion": "95%",
        "proof": "[Chụp thùng rác], [Chụp khôi phục thành công], [Chụp xóa vĩnh viễn]"
    },
]


function_rows6 = [
    [1, MEMBERS[0], "Đăng nhập, đăng xuất và phân tuyến vai trò - liên quan 5.3.1", "Cài form đăng nhập sinh viên/admin, xử lý phiên đăng nhập, lưu trạng thái đăng nhập và điều hướng theo vai trò.", "Test đăng nhập sinh viên, đăng nhập admin, sai mật khẩu, truy cập trang riêng khi chưa đăng nhập và đăng xuất.", "95%", "[Chụp đăng nhập admin thành công], [Chụp đăng nhập sinh viên thành công], [Chụp bị chuyển về login khi chưa đăng nhập]"],
    [2, MEMBERS[0], "Cấp OTP quên mật khẩu - liên quan 5.3.2", "Cài chức năng nhận tên đăng nhập/email, tạo OTP có thời hạn và gửi hướng dẫn đặt lại mật khẩu.", "Test gửi OTP bằng username, gửi OTP bằng email, tài khoản không tồn tại và trường định danh bỏ trống.", "95%", "[Chụp gửi OTP thành công], [Chụp lỗi tài khoản không tồn tại]"],
    [3, MEMBERS[0], "Đặt lại mật khẩu bằng OTP - liên quan 5.3.3", "Cài form nhập OTP, mật khẩu mới, xác nhận mật khẩu và cập nhật mật khẩu khi OTP hợp lệ.", "Test OTP đúng, OTP sai, OTP hết hạn, mật khẩu xác nhận không khớp và mật khẩu mới quá ngắn.", "95%", "[Chụp đặt lại mật khẩu thành công], [Chụp lỗi OTP], [Chụp lỗi xác nhận mật khẩu]"],
    [4, MEMBERS[0], "Thống kê dashboard quản trị - liên quan 5.3.4", "Cài tổng hợp số liệu dashboard, thẻ chỉ số, biểu đồ doanh thu, biểu đồ đăng ký và bảng sinh viên còn nợ.", "Test hiển thị thẻ số liệu, biểu đồ doanh thu, biểu đồ đăng ký, bảng sinh viên nợ và trạng thái không có dữ liệu.", "95%", "[Chụp dashboard admin], [Chụp biểu đồ doanh thu], [Chụp bảng sinh viên nợ]"],
    [5, MEMBERS[0], "Quản lý tỉnh/thành phố - liên quan 5.3.5", "Cài thêm, sửa, xóa, lọc, tìm kiếm, xem chi tiết tỉnh/thành phố và kiểm tra dữ liệu tham chiếu.", "Test thêm tỉnh, sửa tỉnh, xóa tỉnh, tìm theo loại, lọc trạng thái, xem chi tiết và lỗi tỉnh đang có phường/xã.", "95%", "[Chụp thêm tỉnh thành công], [Chụp lỗi xóa tỉnh đang được dùng]"],
    [6, MEMBERS[0], "Quản lý phường/xã - liên quan 5.3.6", "Cài thêm, sửa, xóa, lọc theo tỉnh, loại, khu vực, trạng thái và xem chi tiết phường/xã.", "Test thêm phường/xã, sửa phường/xã, xóa phường/xã, lọc theo tỉnh, lọc khu vực và lỗi thiếu tỉnh.", "95%", "[Chụp thêm phường xã thành công], [Chụp lọc theo tỉnh], [Chụp lỗi thiếu tỉnh]"],
    [7, MEMBERS[0], "Quản lý năm học - liên quan 5.3.7", "Cài danh mục năm học, tự đồng bộ năm bắt đầu/kết thúc theo mã và kiểm tra ràng buộc thời gian.", "Test thêm năm học, sửa năm học, xóa năm học, tìm kiếm, lọc trạng thái và lỗi năm kết thúc nhỏ hơn năm bắt đầu.", "95%", "[Chụp thêm năm học], [Chụp lỗi năm học không hợp lệ]"],
    [8, MEMBERS[0], "Quản lý học kỳ và mốc nghiệp vụ - liên quan 5.3.8", "Cài học kỳ, mốc đăng ký, mốc cứu xét, mốc thu học phí, trạng thái chốt đăng ký và trạng thái mở thu.", "Test thêm học kỳ, sửa học kỳ, lọc trạng thái, lọc chốt đăng ký, lọc mở thu và lỗi thứ tự ngày không hợp lệ.", "95%", "[Chụp thêm học kỳ], [Chụp lỗi mốc thời gian], [Chụp lọc mở thu]"],
    [9, MEMBERS[0], "Quản lý tiết học - liên quan 5.3.9", "Cài danh mục tiết học với giờ bắt đầu, giờ kết thúc, thứ tự, trạng thái và ràng buộc lịch học.", "Test thêm tiết, sửa tiết, xóa tiết, tìm theo mã/tên, lỗi giờ kết thúc trước giờ bắt đầu và lỗi tiết đang được lớp dùng.", "95%", "[Chụp thêm tiết học], [Chụp lỗi giờ tiết học]"],
    [10, MEMBERS[0], "Quản lý khoa - liên quan 5.3.10", "Cài danh mục khoa, thông tin liên hệ, trưởng khoa, số ngành, số môn và xem chi tiết khoa.", "Test thêm khoa, sửa khoa, xóa khoa, tìm theo viết tắt, lỗi email và lỗi khoa đang được ngành/môn tham chiếu.", "95%", "[Chụp thêm khoa], [Chụp lỗi email khoa], [Chụp chi tiết khoa]"],
    [11, MEMBERS[0], "Quản lý ngành học - liên quan 5.3.11", "Cài danh mục ngành học theo khoa, số tín chỉ tối thiểu, thời gian đào tạo và trạng thái.", "Test thêm ngành, sửa ngành, xóa ngành, tìm theo khoa, lỗi trùng mã ngành và xem thông tin chương trình đào tạo theo ngành.", "95%", "[Chụp thêm ngành], [Chụp lỗi trùng ngành]"],
    [12, MEMBERS[0], "Tìm kiếm, lọc, phân trang và xem chi tiết danh mục nền - liên quan 5.3.5 đến 5.3.11", "Cài quy ước toolbar chung cho các danh mục nền, lưu tiêu chí tìm kiếm, lọc và mở popup chi tiết bản ghi.", "Test tìm kiếm từng tiêu chí, lọc trạng thái, đổi trang, click dòng xem chi tiết và reset bộ lọc.", "95%", "[Chụp tìm kiếm danh mục], [Chụp popup chi tiết], [Chụp phân trang]"],
    [13, MEMBERS[0], "Kiểm tra ràng buộc khi xóa danh mục nền - liên quan 5.3.5 đến 5.3.11", "Cài kiểm tra bản ghi đang được bảng khác sử dụng trước khi xóa mềm danh mục địa danh, năm học, học kỳ, tiết, khoa, ngành.", "Test xóa bản ghi chưa dùng, xóa bản ghi đang được tham chiếu, thông báo lỗi ràng buộc và dữ liệu còn nguyên sau lỗi.", "95%", "[Chụp xóa thành công], [Chụp lỗi dữ liệu đang được tham chiếu]"],
    [14, MEMBERS[1], "Quản lý hồ sơ sinh viên - liên quan 5.3.12", "Cài thêm, sửa, xóa sinh viên, kiểm tra MSSV, CCCD, email, số điện thoại, ngành, dân tộc, địa chỉ và đối tượng ưu tiên.", "Test thêm sinh viên, sửa sinh viên, xóa sinh viên, lỗi trùng MSSV, lỗi trùng CCCD và tìm kiếm sinh viên.", "95%", "[Chụp thêm sinh viên], [Chụp lỗi trùng MSSV], [Chụp tìm kiếm sinh viên]"],
    [15, MEMBERS[1], "Import/export và ảnh đại diện sinh viên - liên quan 5.3.12", "Cài nhập Excel sinh viên, xuất Excel theo bộ lọc và cập nhật ảnh đại diện cho từng sinh viên.", "Test import đúng mẫu, import sai mẫu, lỗi từng dòng, export Excel, upload ảnh hợp lệ và lỗi file ảnh không hợp lệ.", "95%", "[Chụp import sinh viên], [Chụp lỗi import], [Chụp upload ảnh sinh viên]"],
    [16, MEMBERS[1], "Quản lý môn học và tự tính tín chỉ - liên quan 5.3.13", "Cài thêm, sửa, xóa môn học, khóa mã môn, tính số tín chỉ theo số tiết và loại môn.", "Test thêm môn lý thuyết, thêm môn thực hành, sửa số tiết, kiểm tra tín chỉ tự tính, xóa môn và lỗi môn đang được dùng.", "95%", "[Chụp thêm môn học], [Chụp tín chỉ tự tính], [Chụp lỗi xóa môn]"],
    [17, MEMBERS[1], "Import/export môn học - liên quan 5.3.13", "Cài nhập Excel môn học, xuất Excel danh mục môn theo bộ lọc và hiển thị kết quả import theo từng dòng.", "Test import thành công, import thiếu cột, import trùng mã môn, lỗi loại môn không hợp lệ và export Excel.", "95%", "[Chụp import môn học thành công], [Chụp lỗi import môn học], [Chụp export môn học]"],
    [18, MEMBERS[1], "Mở môn học theo học kỳ - liên quan 5.3.14", "Cài chức năng chọn học kỳ, chọn môn bằng popup, thêm/sửa/xóa môn học mở và lọc theo khoa, trạng thái.", "Test thêm môn mở, sửa trạng thái, xóa môn mở, popup chọn môn, lỗi trùng môn trong học kỳ và lọc học kỳ.", "95%", "[Chụp thêm môn mở], [Chụp popup chọn môn], [Chụp lỗi trùng môn mở]"],
    [19, MEMBERS[1], "Quản lý lớp học và lịch học - liên quan 5.3.15", "Cài thêm, sửa, xóa lớp, chọn môn, giảng viên, phòng, thứ, tiết và kiểm tra trùng lịch.", "Test thêm lớp, sửa lịch, xóa lớp, lỗi trùng phòng, lỗi trùng giảng viên, lọc trạng thái mở và sắp xếp sĩ số.", "95%", "[Chụp thêm lớp], [Chụp lỗi trùng lịch], [Chụp lọc lớp]"],
    [20, MEMBERS[1], "Mở/đóng lớp và xem sinh viên trong lớp - liên quan 5.3.15", "Cài mở lớp theo học kỳ, đóng lớp mở, cập nhật sĩ số và xem danh sách sinh viên đăng ký trong lớp.", "Test mở lớp thành công, đóng lớp, lỗi đóng lớp đang có đăng ký không hợp lệ, xem danh sách sinh viên và đổi học kỳ.", "90%", "[Chụp mở lớp], [Chụp đóng lớp], [Chụp danh sách sinh viên lớp]"],
    [21, MEMBERS[1], "Quản lý phòng học theo học kỳ - liên quan 5.3.16", "Cài phòng học, học kỳ áp dụng, sức chứa, loại phòng, tình trạng dùng và popup lớp đang dùng phòng.", "Test thêm phòng, sửa phòng, xóa phòng, lọc học kỳ, lọc phòng trống/đang dùng và xem lớp đang dùng.", "95%", "[Chụp thêm phòng], [Chụp lớp đang dùng], [Chụp lỗi xóa phòng]"],
    [22, MEMBERS[1], "Quản lý giảng viên theo học kỳ - liên quan 5.3.17", "Cài giảng viên, học hàm, học vị, khoa, email, học kỳ áp dụng và popup lớp giảng dạy.", "Test thêm giảng viên, sửa giảng viên, lọc học kỳ, xem lớp giảng dạy, lỗi email và lỗi xóa giảng viên đang có lớp.", "95%", "[Chụp thêm giảng viên], [Chụp lớp giảng dạy], [Chụp lỗi email]"],
    [23, MEMBERS[1], "Quản lý ràng buộc môn học - liên quan 5.3.18", "Cài môn tiên quyết/học trước, chọn môn bằng popup, kiểm tra trùng ràng buộc và trạng thái áp dụng.", "Test thêm ràng buộc, sửa ràng buộc, xóa ràng buộc, lỗi tự phụ thuộc, lỗi trùng ràng buộc và lọc loại điều kiện.", "95%", "[Chụp thêm ràng buộc], [Chụp lỗi tự phụ thuộc], [Chụp popup chọn môn]"],
    [24, MEMBERS[1], "Quản lý chương trình học - liên quan 5.3.19", "Cài môn trong chương trình đào tạo theo ngành, học kỳ dự kiến, popup chọn môn và nhóm môn theo học kỳ.", "Test thêm môn vào CTĐT, sửa học kỳ dự kiến, xóa môn khỏi CTĐT, lỗi trùng môn trong ngành và lọc theo ngành.", "95%", "[Chụp thêm môn CTĐT], [Chụp lỗi trùng CTĐT], [Chụp lọc CTĐT]"],
    [25, MEMBERS[1], "Quản lý môn đã học và import kết quả - liên quan 5.3.20", "Cài thêm/sửa/xóa môn đã học, xem chi tiết theo sinh viên và import Excel kết quả học tập.", "Test thêm kết quả, sửa kết quả, xóa kết quả, import đúng mẫu, lỗi trùng lần học và lọc theo học kỳ/kết quả.", "95%", "[Chụp thêm môn đã học], [Chụp import kết quả], [Chụp lỗi trùng lần học]"],
    [26, MEMBERS[1], "Quản lý đăng ký và đơn cứu xét phía admin - liên quan 5.3.21, 5.3.22", "Cài tra cứu phiếu đăng ký, xem chi tiết đúng học kỳ, chốt/hủy chốt đăng ký, duyệt/từ chối đơn cứu xét.", "Test xem chi tiết phiếu, chốt đăng ký, lỗi chốt khi còn đơn chờ, duyệt đơn, từ chối đơn và lỗi duyệt do trùng lịch.", "90%", "[Chụp chi tiết đăng ký], [Chụp chốt đăng ký], [Chụp duyệt đơn cứu xét]"],
    [27, MEMBERS[2], "Dashboard sinh viên - liên quan 5.3.23", "Cài tổng quan tín chỉ hoàn thành, công nợ, lịch học hôm nay, thông báo mới và lối tắt chức năng.", "Test tải dashboard, hiển thị tín chỉ, hiển thị công nợ, lịch hôm nay, thông báo mới và trạng thái không có dữ liệu.", "95%", "[Chụp dashboard sinh viên], [Chụp thẻ công nợ], [Chụp thông báo mới]"],
    [28, MEMBERS[2], "Tra cứu học phần mở và lọc lịch - liên quan 5.3.24", "Cài tìm kiếm học phần, lọc loại học phần, thứ, tiết bắt đầu, tiết kết thúc và hiển thị học phần có thể đăng ký.", "Test tìm theo mã môn, tìm theo tên môn, lọc loại học phần, lọc thứ/tiết và trạng thái hết thời gian đăng ký.", "95%", "[Chụp tra cứu học phần], [Chụp lọc theo tiết], [Chụp hết hạn đăng ký]"],
    [29, MEMBERS[2], "Đăng ký học phần - liên quan 5.3.24", "Cài đăng ký lớp học phần, kiểm tra cửa sổ đăng ký, sĩ số, trùng môn, trùng lịch, tiên quyết và giới hạn tín chỉ.", "Test đăng ký thành công, lỗi trùng môn, lỗi trùng lịch, lỗi hết chỗ, lỗi chưa đạt tiên quyết và cập nhật phiếu đăng ký.", "90%", "[Chụp đăng ký thành công], [Chụp lỗi trùng lịch], [Chụp lỗi tiên quyết]"],
    [30, MEMBERS[2], "Gửi và theo dõi đơn cứu xét sinh viên - liên quan 5.3.24, 5.3.25", "Cài gửi đơn thêm, đổi, hủy học phần; theo dõi trạng thái đơn và hủy đơn khi còn chờ duyệt.", "Test gửi đơn thêm, gửi đơn đổi lớp, gửi đơn hủy lớp, hủy đơn chờ duyệt và xem đơn đã được duyệt/từ chối.", "90%", "[Chụp gửi đơn cứu xét], [Chụp hủy đơn], [Chụp đơn bị từ chối]"],
    [31, MEMBERS[2], "Xem phiếu đăng ký và hủy học phần - liên quan 5.3.25", "Cài xem phiếu đăng ký theo năm học/học kỳ, trạng thái học phần, xác nhận hủy và cập nhật danh sách sau hủy.", "Test lọc phiếu theo học kỳ, xem môn đăng ký, hủy học phần thành công, hủy khi quá hạn và modal xác nhận.", "90%", "[Chụp phiếu đăng ký], [Chụp hủy học phần], [Chụp lỗi hủy quá hạn]"],
    [32, MEMBERS[2], "Xem môn đã học của sinh viên - liên quan 5.3.26", "Cài lịch sử môn đã học, lọc loại môn, khoa, học kỳ, kết quả và tổng hợp môn qua/rớt.", "Test tìm môn đã học, lọc học kỳ, lọc qua môn, lọc rớt, lọc khoa và trạng thái không có dữ liệu.", "95%", "[Chụp môn đã học], [Chụp lọc kết quả], [Chụp không có dữ liệu]"],
    [33, MEMBERS[2], "Xem chương trình đào tạo - liên quan 5.3.27", "Cài hiển thị chương trình đào tạo theo ngành của sinh viên, tiến độ tín chỉ và lọc môn theo trạng thái.", "Test tải CTĐT, lọc môn đã học, lọc môn chưa học, tính tiến độ tín chỉ và trường hợp ngành chưa có CTĐT.", "95%", "[Chụp CTĐT], [Chụp tiến độ tích lũy], [Chụp lọc môn chưa học]"],
    [34, MEMBERS[2], "Xem thời khóa biểu - liên quan 5.3.28", "Cài lịch học trong tuần từ các lớp đã đăng ký, nhóm theo thứ, sắp xếp tiết và hiển thị phòng/giảng viên.", "Test tải lịch học, đổi học kỳ, hiển thị phòng học, sắp xếp theo tiết và trạng thái chưa có lịch.", "95%", "[Chụp thời khóa biểu], [Chụp đổi học kỳ], [Chụp chưa có lịch]"],
    [35, MEMBERS[2], "Xem học phí và công nợ cá nhân - liên quan 5.3.29", "Cài thẻ tổng học phí, đã đóng, còn lại, bảng học phí theo học kỳ và chi tiết học phí.", "Test tải học phí, xem chi tiết học phí, trạng thái chưa đóng, đóng một phần, đã đóng đủ và quá hạn.", "95%", "[Chụp học phí sinh viên], [Chụp chi tiết học phí], [Chụp trạng thái quá hạn]"],
    [36, MEMBERS[2], "Tạo thanh toán học phí - liên quan 5.3.29", "Cài chọn phiếu thu còn nợ, nhập số tiền, chọn loại thanh toán, phương thức và tạo giao dịch thanh toán.", "Test thanh toán tiền mặt, thanh toán trực tuyến, lỗi số tiền nhỏ hơn mức tối thiểu, lỗi số tiền vượt còn nợ và thanh toán thất bại.", "90%", "[Chụp tạo thanh toán], [Chụp thanh toán thành công], [Chụp thanh toán thất bại]"],
    [37, MEMBERS[2], "Xem phiếu thu sinh viên - liên quan 5.3.30", "Cài danh sách phiếu thu của sinh viên theo học kỳ, trạng thái xác nhận, mã giao dịch và modal chi tiết.", "Test lọc phiếu thu theo học kỳ, xem chi tiết phiếu, trạng thái thành công, trạng thái thất bại, trạng thái hủy và chưa có phiếu.", "95%", "[Chụp phiếu thu sinh viên], [Chụp chi tiết phiếu thu], [Chụp trạng thái thất bại]"],
    [38, MEMBERS[2], "Xem và đánh dấu thông báo sinh viên - liên quan 5.3.31", "Cài danh sách thông báo, lọc đã đọc/chưa đọc, mở chi tiết thông báo và cập nhật số thông báo chưa đọc.", "Test xem danh sách thông báo, lọc chưa đọc, mở chi tiết, đánh dấu đã đọc và cập nhật bộ đếm.", "95%", "[Chụp thông báo sinh viên], [Chụp chi tiết thông báo], [Chụp lọc chưa đọc]"],
    [39, MEMBERS[2], "Cập nhật hồ sơ cá nhân sinh viên và quản trị viên - liên quan 5.3.32, 5.3.33", "Cài hiển thị hồ sơ, trường khóa chỉ xem, cập nhật thông tin được phép sửa và cập nhật ảnh đại diện.", "Test lưu số điện thoại, lưu địa chỉ, lưu hồ sơ admin, upload ảnh, lỗi ảnh không hợp lệ và cảnh báo trường không được sửa.", "95%", "[Chụp hồ sơ sinh viên], [Chụp hồ sơ admin], [Chụp cảnh báo trường khóa]"],
    [40, MEMBERS[3], "Tra cứu công nợ học phí admin - liên quan 5.3.34", "Cài tìm kiếm MSSV/họ tên, lọc học kỳ, lọc trạng thái nợ và bảng công nợ học phí.", "Test tìm theo MSSV, tìm theo họ tên, lọc học kỳ, lọc chưa đóng, lọc quá hạn và lọc đã đóng đủ.", "95%", "[Chụp công nợ học phí], [Chụp lọc quá hạn], [Chụp tìm MSSV]"],
    [41, MEMBERS[3], "Xem chi tiết học phí và trạng thái nợ - liên quan 5.3.34", "Cài popup chi tiết học phí, tổng phải đóng, đã đóng, còn nợ, hạn đóng và danh sách học phần tạo công nợ.", "Test mở chi tiết, kiểm tra tổng tiền, kiểm tra miễn giảm, kiểm tra trạng thái quá hạn và dữ liệu không phát sinh học phí.", "95%", "[Chụp chi tiết học phí], [Chụp miễn giảm], [Chụp chưa phát sinh học phí]"],
    [42, MEMBERS[3], "Lập phiếu thu đơn lẻ và hàng loạt - liên quan 5.3.35", "Cài lập phiếu thu theo sinh viên, học kỳ, số tiền, phương thức và tạo hàng loạt phiếu thu cho các đăng ký còn nợ.", "Test lập phiếu đơn lẻ, tạo hàng loạt theo học kỳ, lỗi sinh viên không có nợ, lỗi phiếu trùng hiệu lực và cập nhật danh sách.", "90%", "[Chụp lập phiếu thu], [Chụp tạo hàng loạt], [Chụp lỗi không có nợ]"],
    [43, MEMBERS[3], "Xử lý trạng thái phiếu thu và xuất Excel - liên quan 5.3.35", "Cài xác nhận phiếu thu, đánh dấu thất bại, hủy phiếu, in phiếu và xuất Excel danh sách phiếu thu.", "Test xác nhận thành công, đánh dấu thất bại, hủy phiếu, in phiếu, xuất Excel và kiểm tra công nợ cập nhật.", "90%", "[Chụp xác nhận phiếu], [Chụp phiếu thất bại], [Chụp export phiếu thu]"],
    [44, MEMBERS[3], "Quản lý đơn giá tín chỉ - liên quan 5.3.36", "Cài đơn giá theo loại môn, loại học, học kỳ áp dụng, trạng thái và ghi chú.", "Test thêm đơn giá, sửa đơn giá, xóa đơn giá, lỗi trùng loại giá trong học kỳ, lỗi đơn giá âm và lọc học kỳ.", "95%", "[Chụp thêm đơn giá], [Chụp lỗi trùng đơn giá], [Chụp lọc đơn giá]"],
    [45, MEMBERS[3], "Quản lý đối tượng ưu tiên và sinh viên ưu tiên - liên quan 5.3.37", "Cài đối tượng ưu tiên, tỷ lệ giảm, độ ưu tiên tự tính, thêm sinh viên vào đối tượng và import danh sách sinh viên.", "Test thêm đối tượng, sửa tỷ lệ giảm, độ ưu tiên cập nhật, thêm SV, import SV, lỗi MSSV trùng và xóa SV khỏi đối tượng.", "95%", "[Chụp thêm đối tượng], [Chụp thêm SV ưu tiên], [Chụp lỗi MSSV trùng]"],
    [46, MEMBERS[3], "Báo cáo thống kê tổng hợp - liên quan 5.3.38", "Cài thống kê sinh viên, môn học, doanh thu theo tháng, thanh toán theo phương thức và sinh viên theo trạng thái.", "Test báo cáo toàn bộ, báo cáo theo học kỳ, biểu đồ doanh thu, thống kê phương thức thanh toán và dữ liệu rỗng.", "95%", "[Chụp báo cáo thống kê], [Chụp biểu đồ doanh thu], [Chụp lọc báo cáo]"],
    [47, MEMBERS[3], "Báo cáo sinh viên chưa hoàn thành học phí - liên quan 5.3.39", "Cài lọc sinh viên nợ theo học kỳ, khoa, ngành, trạng thái, quá hạn; xuất Excel và in báo cáo.", "Test lọc học kỳ, lọc khoa/ngành, lọc quá hạn, xuất Excel, in báo cáo và kiểm tra tổng còn nợ.", "95%", "[Chụp báo cáo nợ], [Chụp export báo cáo nợ], [Chụp lọc quá hạn]"],
    [48, MEMBERS[3], "Quản lý tài khoản người dùng - liên quan 5.3.40", "Cài tạo tài khoản, xác nhận mật khẩu, lọc vai trò/nhóm, đổi nhóm người dùng và xóa tài khoản.", "Test tạo tài khoản admin, tạo tài khoản sinh viên, lỗi xác nhận mật khẩu, đổi nhóm, xóa tài khoản và tìm theo email.", "90%", "[Chụp tạo tài khoản], [Chụp lỗi xác nhận mật khẩu], [Chụp đổi nhóm]"],
    [49, MEMBERS[3], "Tạo tài khoản hàng loạt và mật khẩu tạm - liên quan 5.3.40", "Cài tạo tài khoản sinh viên theo khoa/ngành, sinh mật khẩu tạm, lưu danh sách tài khoản đã tạo và ẩn/hiện mật khẩu.", "Test tạo hàng loạt theo khoa, tạo hàng loạt theo ngành, sinh mật khẩu, xem danh sách mật khẩu tạm và lỗi sinh viên đã có tài khoản.", "90%", "[Chụp tạo hàng loạt], [Chụp danh sách mật khẩu tạm], [Chụp lỗi đã có tài khoản]"],
    [50, MEMBERS[3], "Phân quyền nhóm và chức năng - liên quan 5.3.41", "Cài nhóm người dùng, quyền truy cập, gán quyền theo nhóm và kiểm soát sidebar/màn hình theo quyền.", "Test thêm nhóm, thêm quyền, gán quyền, xóa quyền, đăng nhập nhóm bị hạn chế và truy cập màn hình không đủ quyền.", "90%", "[Chụp gán quyền], [Chụp sidebar theo quyền], [Chụp lỗi không đủ quyền]"],
    [51, MEMBERS[3], "Quản lý thông báo admin - liên quan 5.3.42", "Cài tạo, sửa, xóa, ghim, xem trước thông báo và gửi theo toàn hệ thống, khoa, ngành hoặc cá nhân.", "Test tạo thông báo chung, tạo theo khoa/ngành, xem trước, ghim, lọc nguồn tự động/thủ công và xóa thông báo.", "95%", "[Chụp tạo thông báo], [Chụp xem trước], [Chụp lọc thông báo]"],
    [52, MEMBERS[3], "Tham số hệ thống và thùng rác - liên quan 5.3.43, 5.3.44", "Cài tham số quy chế đăng ký, ràng buộc Anh văn, điều kiện khóa luận, thùng rác, khôi phục và xóa vĩnh viễn.", "Test cập nhật tham số, lỗi tham số không hợp lệ, khôi phục bản ghi, xóa vĩnh viễn, thao tác hàng loạt và quyền xem loại thùng rác.", "95%", "[Chụp cập nhật tham số], [Chụp khôi phục thùng rác], [Chụp xóa vĩnh viễn]"],
]


def set_run_font(run, size=12, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:cs"), "Times New Roman")


def set_paragraph_text(paragraph, text, size=12, bold=False):
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.05
    return paragraph


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=70, bottom=70, end=70):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_layout(table, widths, font_size=11):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx >= len(row.cells):
                continue
            cell = row.cells[idx]
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    set_run_font(run, size=font_size)


def fill_table(table, data, widths, font_size=11, header_font_size=12):
    for r_idx, row in enumerate(data):
        cells = table.rows[r_idx].cells
        for c_idx, value in enumerate(row):
            p = cells[c_idx].paragraphs[0]
            set_paragraph_text(
                p,
                str(value),
                size=header_font_size if r_idx == 0 else font_size,
                bold=(r_idx == 0),
            )
            if r_idx == 0 or c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if r_idx == 0:
                set_cell_shading(cells[c_idx], "A6A6A6")
    set_table_layout(table, widths, font_size=font_size)


def find_paragraph_element(doc, predicate):
    for p in doc.paragraphs:
        text = " ".join(p.text.split())
        if predicate(text):
            return p._p
    return None


def remove_old_sections(doc):
    body = doc.element.body
    start = find_paragraph_element(
        doc,
        lambda text: text in {"Thiết kế giao diện:", "5. Thiết kế giao diện"}
        or text.startswith("5. Thiết kế giao diện"),
    )
    marker = find_paragraph_element(doc, lambda text: text.startswith("Kết luận:"))
    if start is None or marker is None:
        raise RuntimeError("Không tìm được ranh giới mục 5-6 hoặc phần Kết luận trong DOCX.")

    elements = list(body)
    start_idx = elements.index(start)
    marker_idx = elements.index(marker)
    for element in elements[start_idx:marker_idx]:
        body.remove(element)
    return marker


def add_before(doc, marker, kind, *args, **kwargs):
    if kind == "paragraph":
        obj_ = doc.add_paragraph(*args, **kwargs)
        marker.addprevious(obj_._p)
        return obj_
    if kind == "table":
        obj_ = doc.add_table(*args, **kwargs)
        marker.addprevious(obj_._tbl)
        return obj_
    raise ValueError(kind)


def add_heading(doc, marker, text, level):
    p = add_before(doc, marker, "paragraph", style=f"Heading {level}")
    set_paragraph_text(p, text, size=14 if level <= 2 else 13, bold=(level <= 2))
    return p


def add_normal(doc, marker, text, size=12, bold=False):
    p = add_before(doc, marker, "paragraph")
    set_paragraph_text(p, text, size=size, bold=bold)
    return p


def add_table_before(doc, marker, headers, body_rows, widths, font_size=11, header_font_size=12):
    table = add_before(doc, marker, "table", rows=len(body_rows) + 1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    fill_table(table, [headers] + body_rows, widths, font_size=font_size, header_font_size=header_font_size)
    spacer = add_before(doc, marker, "paragraph")
    set_paragraph_text(spacer, "", size=4)
    return table


def build_document():
    if len(screens) != 44:
        raise RuntimeError(f"Số màn hình phải là 44, hiện có {len(screens)}")
    counts = {member: 0 for member in MEMBERS}
    for screen in screens:
        counts[screen["member"]] += 1
    if any(value != 11 for value in counts.values()):
        raise RuntimeError(f"Phân công không đều: {counts}")
    function_counts = {member: 0 for member in MEMBERS}
    for row in function_rows6:
        function_counts[row[1]] += 1
    if any(value != 13 for value in function_counts.values()):
        raise RuntimeError(f"Phân công chức năng mục 6 không đều: {function_counts}")

    doc = Document(DOCX_PATH)
    marker = remove_old_sections(doc)

    add_heading(doc, marker, "5. Thiết kế giao diện:", 2)
    add_heading(doc, marker, "5.1. Sơ đồ liên kết các màn hình.", 3)
    flow_texts = [
        "Luồng public: người dùng vào / hoặc /login; hệ thống kiểm tra token. Nếu đã đăng nhập thì chuyển quản trị viên về /admin/dashboard và sinh viên về /student/dashboard; nếu chưa đăng nhập thì hiển thị màn hình đăng nhập.",
        "Luồng khôi phục mật khẩu: từ màn hình đăng nhập chọn quên mật khẩu, nhập tên đăng nhập hoặc email tại /forgot-password hoặc /admin/forgot-password, nhận OTP và hoàn tất tại /reset-password.",
        "Luồng quản trị: /admin/dashboard là điểm vào chính. Sidebar admin liên kết đến Địa danh, Đào tạo, Đăng ký, Tài chính, Báo cáo, Hệ thống và Quản trị; mỗi màn hình gọi các route API tương ứng để tìm kiếm, thêm, sửa, xóa, xem chi tiết hoặc xuất dữ liệu.",
        "Luồng sinh viên: /student/dashboard là điểm vào chính. Sidebar sinh viên liên kết đến đăng ký học phần, phiếu đăng ký, môn đã học, thời khóa biểu, chương trình đào tạo, học phí, phiếu thu, hồ sơ và thông báo.",
        "Luồng nghiệp vụ học phí: admin cấu hình đơn giá và đối tượng ưu tiên, sinh viên đăng ký học phần, hệ thống tính công nợ, admin tạo phiếu thu, sinh viên thanh toán, trạng thái phiếu thu cập nhật về công nợ và báo cáo.",
        "[Cần chụp hình sơ đồ liên kết các màn hình].",
    ]
    for text in flow_texts:
        add_normal(doc, marker, text)

    add_heading(doc, marker, "5.2. Danh sách các màn hình:", 3)
    list_rows = []
    for idx, screen in enumerate(screens, 1):
        list_rows.append([
            idx,
            f'{screen["name"]} ({screen["route"]})',
            screen["type"],
            f'{screen["function"]} Ảnh: [Cần chụp hình {screen["name"].lower()}].'
        ])
    add_table_before(
        doc,
        marker,
        ["STT", "Màn hình", "Loại màn hình", "Chức năng"],
        list_rows,
        [0.45, 1.45, 1.15, 3.0],
        font_size=10,
        header_font_size=12,
    )

    add_heading(doc, marker, "5.3. Mô tả các màn hình:", 3)
    for idx, screen in enumerate(screens, 1):
        add_heading(doc, marker, f'5.3.{idx}. {screen["name"]}:', 4)
        add_normal(doc, marker, "a.\tGiao diện")
        add_normal(
            doc,
            marker,
            f'{screen["layout"]} Route: {screen["route"]}. Component/view: {screen["view"]}. Client script: {screen["js"]}. [Cần chụp hình {screen["name"].lower()}].'
        )
        add_normal(doc, marker, "b.\tMô tả các đối tượng trên màn hình:")
        object_rows = [[i + 1] + row for i, row in enumerate(screen["objects"])]
        add_table_before(
            doc,
            marker,
            ["STT", "Tên", "Kiểu", "Ràng buộc", "Chức năng"],
            object_rows,
            [0.55, 1.25, 1.05, 1.75, 1.65],
            font_size=10,
            header_font_size=12,
        )
        add_normal(doc, marker, "c.\tDanh sách biến cố và xử lý tương ứng trên màn hình:")
        event_rows = [
            [i + 1, business_event_label(row[0]), business_flow(screen, row[0], row[1])]
            for i, row in enumerate(screen["events"])
        ]
        add_table_before(
            doc,
            marker,
            ["STT", "Biến cố", "Xử lý"],
            event_rows,
            [0.55, 2.25, 3.45],
            font_size=10,
            header_font_size=12,
        )

    add_heading(doc, marker, "6. Cài đặt và thử nghiệm:", 2)
    add_normal(
        doc,
        marker,
        "Bảng dưới đây liệt kê theo chức năng cài đặt và thử nghiệm, không lấy màn hình làm đơn vị chính. Mỗi chức năng đều ghi rõ màn hình liên quan ở mục 5.3 để đối chiếu với thiết kế giao diện."
    )
    function_table_rows = [
        [
            row[0],
            f'{row[2]}',
            "100%",
            f'Cài đặt: {row[3]} Thử nghiệm: {row[4]}',
        ]
        for row in function_rows6
    ]
    add_table_before(
        doc,
        marker,
        [
            "STT",
            "Chức năng",
            "Mức độ hoàn thành (%)",
            "Ghi chú",
        ],
        function_table_rows,
        [0.55, 1.45, 1.75, 2.5],
        font_size=10,
        header_font_size=12,
    )

    doc.save(OUT_PATH)


if __name__ == "__main__":
    build_document()
