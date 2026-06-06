import os
import re
import zipfile
from dataclasses import dataclass
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont


INPUT = "Nhom5QLTCVTHP.docx"
OUTPUT = os.path.join("outputs", "Nhom5QLTCVTHP_usecase_5_1.docx")
DIAGRAM_DIR = os.path.join("outputs", "usecase_5_1")


@dataclass
class Actor:
    name: str
    x: int
    y: int


@dataclass
class UseCase:
    text: str
    x: int
    y: int
    w: int = 360
    h: int = 96


@dataclass
class DiagramSpec:
    key_text: str
    title: str
    file_name: str
    system_name: str
    actors: list[Actor]
    usecases: list[UseCase]
    links: list[tuple[int, int]]
    notes: list[str]


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        "C:/Windows/Fonts/timesbd.ttf" if bold else "C:/Windows/Fonts/times.ttf",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
    ]
    for path in candidates:
        if os.path.exists(path):
            return ImageFont.truetype(path, size=size)
    return ImageFont.load_default()


FONT_TITLE = font(38, True)
FONT_SYSTEM = font(30, True)
FONT_TEXT = font(27)
FONT_SMALL = font(22)


def wrap_text(draw: ImageDraw.ImageDraw, text: str, max_width: int, fnt) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        test = word if not current else current + " " + word
        bbox = draw.textbbox((0, 0), test, font=fnt)
        if bbox[2] - bbox[0] <= max_width:
            current = test
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_centered_text(draw: ImageDraw.ImageDraw, box, text: str, fnt, fill=(20, 20, 20)):
    x, y, w, h = box
    lines = wrap_text(draw, text, max(10, w - 30), fnt)
    line_heights = []
    total_h = 0
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=fnt)
        lh = bbox[3] - bbox[1] + 4
        line_heights.append(lh)
        total_h += lh
    cy = y + (h - total_h) / 2
    for line, lh in zip(lines, line_heights):
        bbox = draw.textbbox((0, 0), line, font=fnt)
        tw = bbox[2] - bbox[0]
        draw.text((x + (w - tw) / 2, cy), line, font=fnt, fill=fill)
        cy += lh


def draw_actor(draw: ImageDraw.ImageDraw, actor: Actor):
    x, y = actor.x, actor.y
    ink = (34, 34, 34)
    draw.ellipse((x - 24, y - 78, x + 24, y - 30), outline=ink, width=4)
    draw.line((x, y - 30, x, y + 44), fill=ink, width=4)
    draw.line((x - 54, y - 2, x + 54, y - 2), fill=ink, width=4)
    draw.line((x, y + 44, x - 48, y + 108), fill=ink, width=4)
    draw.line((x, y + 44, x + 48, y + 108), fill=ink, width=4)
    draw_centered_text(draw, (x - 125, y + 116, 250, 80), actor.name, FONT_SMALL)


def ellipse_point_towards(uc: UseCase, target: tuple[int, int]) -> tuple[int, int]:
    cx = uc.x + uc.w / 2
    cy = uc.y + uc.h / 2
    tx, ty = target
    dx = tx - cx
    dy = ty - cy
    if abs(dx) / max(1, uc.w / 2) > abs(dy) / max(1, uc.h / 2):
        x = uc.x + (uc.w if dx > 0 else 0)
        y = cy + dy * abs((x - cx) / dx) if dx else cy
    else:
        y = uc.y + (uc.h if dy > 0 else 0)
        x = cx + dx * abs((y - cy) / dy) if dy else cx
    return int(x), int(y)


def actor_anchor(actor: Actor, target: tuple[int, int]) -> tuple[int, int]:
    ax, ay = actor.x, actor.y
    tx, _ = target
    if tx > ax:
        return ax + 60, ay + 10
    return ax - 60, ay + 10


def draw_usecase(draw: ImageDraw.ImageDraw, uc: UseCase):
    fill = (248, 252, 255)
    outline = (25, 84, 140)
    draw.ellipse((uc.x, uc.y, uc.x + uc.w, uc.y + uc.h), fill=fill, outline=outline, width=4)
    draw_centered_text(draw, (uc.x, uc.y, uc.w, uc.h), uc.text, FONT_TEXT)


def draw_diagram(spec: DiagramSpec, output_path: str):
    w, h = 1800, 1080
    im = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(im)
    # Header
    draw.rounded_rectangle((30, 24, w - 30, 92), radius=18, fill=(232, 240, 249), outline=(82, 122, 165), width=2)
    draw_centered_text(draw, (30, 24, w - 60, 68), spec.title, FONT_TITLE)

    # System boundary
    sx, sy, sw, sh = 360, 145, 1080, 780
    draw.rounded_rectangle((sx, sy, sx + sw, sy + sh), radius=26, outline=(45, 90, 135), width=4)
    draw.rectangle((sx + 28, sy - 22, sx + 28 + 500, sy + 24), fill="white")
    draw.text((sx + 42, sy - 19), spec.system_name, font=FONT_SYSTEM, fill=(25, 74, 120))

    for actor in spec.actors:
        draw_actor(draw, actor)
    for uc in spec.usecases:
        draw_usecase(draw, uc)

    for actor_idx, uc_idx in spec.links:
        actor = spec.actors[actor_idx]
        uc = spec.usecases[uc_idx]
        target = (uc.x + uc.w // 2, uc.y + uc.h // 2)
        start = actor_anchor(actor, target)
        end = ellipse_point_towards(uc, (actor.x, actor.y))
        draw.line((start[0], start[1], end[0], end[1]), fill=(70, 70, 70), width=3)

    if spec.notes:
        note_x, note_y = 410, 950
        for idx, note in enumerate(spec.notes):
            draw.text((note_x, note_y + idx * 34), note, font=FONT_SMALL, fill=(70, 70, 70))

    im.save(output_path, "PNG")


def build_specs() -> list[DiagramSpec]:
    return [
        DiagramSpec(
            key_text="Luồng public:",
            title="Sơ đồ use case luồng public và đăng nhập",
            file_name="usecase_01_public_login.png",
            system_name="Hệ thống xác thực và điều hướng",
            actors=[Actor("Người dùng", 170, 460)],
            usecases=[
                UseCase("Truy cập trang chủ /login", 470, 215),
                UseCase("Kiểm tra token đăng nhập", 930, 215),
                UseCase("Đăng nhập sinh viên", 470, 405),
                UseCase("Đăng nhập quản trị viên", 930, 405),
                UseCase("Chuyển dashboard đúng vai trò", 700, 610, 420, 105),
                UseCase("Đăng xuất", 700, 790, 420, 96),
            ],
            links=[(0, 0), (0, 2), (0, 3), (0, 5)],
            notes=[
                "Route liên quan: /, /login, /admin/login, /student/dashboard, /admin/dashboard.",
                "Mục tiêu: xác thực tài khoản và điều hướng đúng vai trò người dùng.",
            ],
        ),
        DiagramSpec(
            key_text="Luồng khôi phục mật khẩu:",
            title="Sơ đồ use case luồng khôi phục mật khẩu",
            file_name="usecase_02_reset_password.png",
            system_name="Hệ thống khôi phục mật khẩu",
            actors=[Actor("Người dùng", 170, 470), Actor("Dịch vụ OTP/Email", 1630, 470)],
            usecases=[
                UseCase("Chọn quên mật khẩu", 485, 205),
                UseCase("Nhập username hoặc email", 485, 390),
                UseCase("Gửi mã OTP", 950, 295),
                UseCase("Xác thực OTP", 485, 575),
                UseCase("Đặt mật khẩu mới", 950, 575),
                UseCase("Quay lại đăng nhập", 725, 770),
            ],
            links=[(0, 0), (0, 1), (0, 3), (0, 4), (1, 2)],
            notes=[
                "Route liên quan: /forgot-password, /admin/forgot-password, /reset-password.",
                "Mục tiêu: cấp OTP, xác thực OTP và cập nhật mật khẩu mới an toàn.",
            ],
        ),
        DiagramSpec(
            key_text="Luồng quản trị:",
            title="Sơ đồ use case luồng quản trị",
            file_name="usecase_03_admin.png",
            system_name="Cổng quản trị hệ thống",
            actors=[Actor("Quản trị viên", 170, 480)],
            usecases=[
                UseCase("Xem dashboard quản trị", 450, 190),
                UseCase("Quản lý địa danh", 910, 190),
                UseCase("Quản lý đào tạo", 450, 350),
                UseCase("Quản lý đăng ký", 910, 350),
                UseCase("Quản lý tài chính", 450, 510),
                UseCase("Xem báo cáo thống kê", 910, 510),
                UseCase("Quản lý hệ thống và phân quyền", 450, 670),
                UseCase("Import/Export dữ liệu", 910, 670),
            ],
            links=[(0, i) for i in range(8)],
            notes=[
                "Điểm vào: /admin/dashboard và sidebar quản trị.",
                "Nhóm chức năng: danh mục, đào tạo, đăng ký, tài chính, báo cáo, hệ thống.",
            ],
        ),
        DiagramSpec(
            key_text="Luồng sinh viên:",
            title="Sơ đồ use case luồng sinh viên",
            file_name="usecase_04_student.png",
            system_name="Cổng sinh viên",
            actors=[Actor("Sinh viên", 170, 480)],
            usecases=[
                UseCase("Xem dashboard sinh viên", 450, 190),
                UseCase("Đăng ký học phần", 910, 190),
                UseCase("Xem phiếu đăng ký", 450, 350),
                UseCase("Tra cứu môn đã học", 910, 350),
                UseCase("Xem thời khóa biểu", 450, 510),
                UseCase("Xem chương trình đào tạo", 910, 510),
                UseCase("Tra cứu học phí và phiếu thu", 450, 670),
                UseCase("Cập nhật hồ sơ, xem thông báo", 910, 670),
            ],
            links=[(0, i) for i in range(8)],
            notes=[
                "Điểm vào: /student/dashboard và sidebar sinh viên.",
                "Mục tiêu: hỗ trợ đăng ký học phần, theo dõi học tập và học phí cá nhân.",
            ],
        ),
        DiagramSpec(
            key_text="Luồng nghiệp vụ học phí:",
            title="Sơ đồ use case luồng nghiệp vụ học phí",
            file_name="usecase_05_tuition.png",
            system_name="Nghiệp vụ học phí",
            actors=[
                Actor("Quản trị viên tài chính", 170, 420),
                Actor("Sinh viên", 170, 750),
                Actor("Cổng thanh toán", 1630, 610),
            ],
            usecases=[
                UseCase("Cấu hình đơn giá tín chỉ", 455, 185),
                UseCase("Cấu hình đối tượng ưu tiên", 920, 185),
                UseCase("Tính công nợ học phí", 690, 365, 420, 102),
                UseCase("Lập phiếu thu", 455, 545),
                UseCase("Thanh toán học phí", 920, 545),
                UseCase("Cập nhật trạng thái phiếu thu", 690, 725, 420, 102),
                UseCase("Báo cáo công nợ và doanh thu", 690, 870, 420, 102),
            ],
            links=[(0, 0), (0, 1), (0, 3), (0, 6), (1, 4), (1, 5), (2, 4), (2, 5)],
            notes=[
                "Luồng dữ liệu: cấu hình giá -> đăng ký học phần -> công nợ -> phiếu thu -> thanh toán -> báo cáo.",
                "Màn hình liên quan: đơn giá tín chỉ, đối tượng ưu tiên, công nợ, phiếu thu, học phí sinh viên, báo cáo.",
            ],
        ),
    ]


def set_run_font(run, size=14, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), "Times New Roman")


def ensure_caption_style(doc: Document):
    if "CaptionHinh" in doc.styles:
        style = doc.styles["CaptionHinh"]
    else:
        style = doc.styles.add_style("CaptionHinh", 1)
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return style


def text_of_paragraph_element(el) -> str:
    return "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()


def insert_paragraph_after(el, doc: Document, style=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    el.addnext(new_p)
    paragraph = Paragraph(new_p, doc._body)
    if style is not None:
        paragraph.style = style
    return paragraph


def remove_paragraph_element(el):
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def remove_existing_usecase_diagrams(doc: Document):
    body = list(doc.element.body)
    for idx, el in enumerate(body):
        if el.tag != qn("w:p"):
            continue
        txt = text_of_paragraph_element(el)
        if "Sơ đồ use case luồng" not in txt:
            continue
        # Caption follows the generated image paragraph. Remove image paragraph immediately before it.
        prev = body[idx - 1] if idx > 0 else None
        remove_paragraph_element(el)
        if prev is not None and prev.tag == qn("w:p") and prev.findall(".//" + qn("a:blip")):
            remove_paragraph_element(prev)


def remove_placeholder(doc: Document):
    for el in list(doc.element.body):
        if el.tag == qn("w:p") and "[Cần chụp hình sơ đồ liên kết các màn hình]" in text_of_paragraph_element(el):
            remove_paragraph_element(el)


def max_existing_figure_number(doc: Document) -> int:
    max_no = 0
    for p in doc.paragraphs:
        txt = p.text.strip()
        m = re.match(r"^Hình\s+(\d+):", txt)
        if m:
            max_no = max(max_no, int(m.group(1)))
    return max_no


def add_diagrams_to_docx(doc: Document, specs: Iterable[DiagramSpec]):
    caption_style = ensure_caption_style(doc)
    start_no = max_existing_figure_number(doc) + 1
    specs = list(specs)
    figure_numbers = {spec.key_text: start_no + idx for idx, spec in enumerate(specs)}
    body = list(doc.element.body)
    by_key: dict[str, object] = {}
    for el in body:
        if el.tag != qn("w:p"):
            continue
        text = text_of_paragraph_element(el)
        for spec in specs:
            if text.startswith(spec.key_text):
                by_key[spec.key_text] = el

    missing = [spec.key_text for spec in specs if spec.key_text not in by_key]
    if missing:
        raise RuntimeError(f"Không tìm thấy đoạn luồng để chèn sơ đồ: {missing}")

    inserted = []
    # Insert in reverse body order so added paragraphs do not affect later insertion anchors.
    ordered_specs = sorted(specs, key=lambda s: body.index(by_key[s.key_text]), reverse=True)
    for spec in ordered_specs:
        anchor = by_key[spec.key_text]
        image_path = os.path.join(DIAGRAM_DIR, spec.file_name)

        image_p = insert_paragraph_after(anchor, doc)
        image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_p.paragraph_format.space_before = Pt(6)
        image_p.paragraph_format.space_after = Pt(3)
        image_p.add_run().add_picture(image_path, width=Inches(6.35))

        caption_p = insert_paragraph_after(image_p._element, doc, caption_style)
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_p.paragraph_format.space_before = Pt(0)
        caption_p.paragraph_format.space_after = Pt(8)
        run = caption_p.add_run(f"Hình {figure_numbers[spec.key_text]}: {spec.title}")
        set_run_font(run, 14)
        inserted.append(caption_p.text)
    return list(reversed(inserted))


def main():
    if not zipfile.is_zipfile(INPUT):
        raise SystemExit(f"{INPUT} không phải DOCX hợp lệ")

    os.makedirs(DIAGRAM_DIR, exist_ok=True)
    specs = build_specs()
    for spec in specs:
        draw_diagram(spec, os.path.join(DIAGRAM_DIR, spec.file_name))

    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
    doc = Document(INPUT)
    remove_existing_usecase_diagrams(doc)
    remove_placeholder(doc)
    inserted = add_diagrams_to_docx(doc, specs)
    doc.save(OUTPUT)

    print(f"saved={OUTPUT}")
    print(f"diagrams={len(specs)}")
    for item in inserted:
        print(item)


if __name__ == "__main__":
    main()
